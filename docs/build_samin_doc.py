#!/usr/bin/env python3
"""
Build a DOCX file explaining ONLY Samin's modules (M2 + M3) in simple terms.
Includes connections to teammates' modules and how they attach.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ---- Style helpers ---------------------------------------------------------

def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_bullet(text, level=0):
    return doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Bullet 2')

def add_number(text):
    return doc.add_paragraph(text, style='List Number')

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shading = p.paragraph_format.element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5'
    })
    shading.append(shd)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

def add_page_break():
    doc.add_page_break()

# ============================================================================
# TITLE PAGE
# ============================================================================

title = doc.add_heading('', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CareerBridge\nMy Modules: M2 Job Management & M3 Application Tracking')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('A simple, technical explanation of my code and how it connects to the rest of the project')
run.font.size = Pt(13)
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nSamin Sarwat (A22EC4040)\nSECJ3483 Web Technology - Group Project - Team CHAMPION\n')
run.font.size = Pt(12)

add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================

add_heading('Table of Contents', level=1)

contents = [
    '1. What My Modules Do (In Simple Words)',
    '2. Where My Code Lives (File Locations)',
    '3. How My Modules Connect to Teammates\' Modules',
    '4. The Database Tables I Designed',
    '5. My API Endpoints (What URLs My Code Handles)',
    '6. Module 2: Job Management - Backend Code Explained',
    '    6a. Listing Jobs (GET /api/jobs) - with search and filter',
    '    6b. Viewing a Single Job (GET /api/jobs/{id})',
    '    6c. Creating a Job (POST /api/jobs) - admin only',
    '    6d. Updating a Job (PUT /api/jobs/{id}) - admin only',
    '    6e. Deleting a Job (DELETE /api/jobs/{id}) - soft delete',
    '    6f. Input Validation (Backend)',
    '7. Module 3: Application Tracking - Backend Code Explained',
    '    7a. Listing Applications (GET /api/applications) - role-aware',
    '    7b. Applying for a Job (POST /api/applications) - student only',
    '    7c. Changing Application Status (PUT /api/applications/{id}/status) - admin only',
    '    7d. Withdrawing an Application (DELETE /api/applications/{id}) - student only',
    '8. My Frontend Code (Vue.js Views)',
    '    8a. BrowseJobs.vue - student browses and applies',
    '    8b. JobDetail.vue - student views full job info',
    '    8c. ManageJobs.vue - admin creates, edits, deactivates jobs',
    '    8d. MyApplications.vue - student tracks and withdraws',
    '    8e. ManageApplications.vue - admin reviews and changes status',
    '9. Shared Code I Use (Built by Teammates)',
    '10. Security in My Modules',
    '11. Database Demo Commands (If DR Asks to See the Database)',
    '12. Questions the Lecturer Might Ask (With Answers)',
]

for item in contents:
    add_para(item, size=11)

add_page_break()

# ============================================================================
# 1. WHAT MY MODULES DO
# ============================================================================

add_heading('1. What My Modules Do (In Simple Words)', level=1)

add_heading('Module 2 (M2): Job Management', level=2)
add_para(
    'This module is like a job board. It lets admins post job listings '
    'and students browse them. Think of it as the "jobs section" of the website.'
)
add_para('What admins can do:', bold=True)
add_bullet('Create a new job post (title, company, type, deadline, description, requirements)')
add_bullet('Edit an existing job post')
add_bullet('Deactivate a job (soft delete - hides it from students but keeps the data in the database)')
add_para('What students can do:', bold=True)
add_bullet('Browse all active jobs on a grid page')
add_bullet('Search jobs by title or company name')
add_bullet('Filter jobs by type (internship, full-time, part-time)')
add_bullet('Click a job to see its full details on a separate page')

add_para('')
add_heading('Module 3 (M3): Application Tracking', level=2)
add_para(
    'This module is like an application tracker. It lets students apply for jobs '
    'and admins track and manage those applications. Think of it as the '
    '"application system" that works on top of the job board.'
)
add_para('What students can do:', bold=True)
add_bullet('Apply to a job with an optional cover letter (max 5000 characters)')
add_bullet('See all their applications in one place with status badges')
add_bullet('Filter their applications by status (pending, reviewed, accepted, rejected)')
add_bullet('Withdraw an application - but only if it is still "pending"')
add_para('What admins can do:', bold=True)
add_bullet('See every application in the system (from all students)')
add_bullet('Change the status of any application (pending -> reviewed -> accepted/rejected)')
add_bullet('Filter applications by status using tabs')

add_para('')
add_para('Important rule: A student cannot apply to the same job twice. '
         'The system blocks this both in the code and in the database.', bold=True)

add_page_break()

# ============================================================================
# 2. WHERE MY CODE LIVES
# ============================================================================

add_heading('2. Where My Code Lives (File Locations)', level=1)

add_para('If the lecturer asks "where is your code?", here is the complete list of files I wrote:')

add_para('')
add_para('Backend (PHP) - 2 controller files:', bold=True)
add_code('''backend/src/Modules/Jobs/JobController.php
    -> M2: Full CRUD for job listings
    -> Methods: index(), show(), create(), update(), delete(), validate()

backend/src/Modules/Applications/ApplicationController.php
    -> M3: Apply, list, update status, withdraw
    -> Methods: index(), create(), updateStatus(), delete(), readJson()''')

add_para('')
add_para('Frontend (Vue.js) - 5 view files:', bold=True)
add_code('''frontend/src/views/BrowseJobs.vue
    -> M2: Student browses jobs, searches, filters, and applies

frontend/src/views/JobDetail.vue
    -> M2: Student views full details of a single job

frontend/src/views/ManageJobs.vue
    -> M2: Admin creates, edits, and deactivates jobs

frontend/src/views/MyApplications.vue
    -> M3: Student tracks their applications and withdraws

frontend/src/views/ManageApplications.vue
    -> M3: Admin reviews applications and changes their status''')

add_para('')
add_para('Database - 2 table definitions in schema.sql:', bold=True)
add_code('''database/schema.sql
    -> I designed the `jobs` table (columns, foreign keys, ENUM types)
    -> I designed the `applications` table (columns, foreign keys, unique constraint)

database/seed.sql
    -> I wrote the seed data for jobs (10 job listings)
    -> I wrote the seed data for applications (5 applications with different statuses)''')

add_para('')
add_para('Total: 2 backend controllers + 5 frontend views + 2 database tables = MY CODE', bold=True)

add_page_break()

# ============================================================================
# 3. HOW MY MODULES CONNECT TO TEAMMATES
# ============================================================================

add_heading('3. How My Modules Connect to Teammates\' Modules', level=1)

add_para(
    'My modules do not work alone. They connect to other modules built by my teammates. '
    'Here is every connection and how it works:'
)

add_para('')

# Connection 1: Jobs -> Companies
add_heading('Connection 1: My Jobs table links to Mariam\'s Companies table (M4)', level=2)
add_para('How they connect:', bold=True)
add_para(
    'Every job belongs to a company. My `jobs` table has a column called `company_id` '
    'which is a foreign key that points to the `companies` table (designed by Mariam). '
    'This means you cannot create a job without choosing a company first.'
)
add_para('Where you can see this connection in my code:', bold=True)
add_bullet('In JobController.php index() - I use JOIN to get company name and location along with job data')
add_bullet('In JobController.php show() - I JOIN with companies to show company info on the job detail page')
add_bullet('In JobController.php create() - the admin must select a company_id from a dropdown')
add_bullet('In ManageJobs.vue - the form loads companies from Mariam\'s API (GET /api/companies) to fill the dropdown')
add_para('Example SQL that shows the connection:', bold=True)
add_code('''SELECT j.title, c.name AS company_name, c.location
FROM jobs j
JOIN companies c ON c.company_id = j.company_id
WHERE j.is_active = 1;''')
add_para(
    'In simple words: My job posts need a company to belong to. '
    'Mariam built the company management, and I link my jobs to her companies.'
)

add_para('')

# Connection 2: Jobs -> Users (Auth)
add_heading('Connection 2: My Jobs table links to Areeb\'s Users table (M1)', level=2)
add_para('How they connect:', bold=True)
add_para(
    'Every job has a `posted_by` column which is a foreign key to the `users` table. '
    'This records which admin created the job. The user_id comes from the JWT token '
    '(Areeb\'s authentication system), not from user input. This means the system '
    'always knows who posted each job, and nobody can fake it.'
)
add_para('Where you can see this connection in my code:', bold=True)
add_bullet('In JobController.php create() - $payload[\'user_id\'] comes from the JWT token')
add_bullet('In JobController.php show() - I JOIN with users to show "posted_by_name"')
add_bullet('In index.php - my POST/PUT/DELETE routes use Areeb\'s JwtMiddleware to check admin role')
add_para(
    'In simple words: Areeb built the login system. My job posts record who created them '
    'using the user ID from the login token. Only admins can create/edit/delete jobs - '
    'this is enforced by Areeb\'s JWT middleware on my routes.'
)

add_para('')

# Connection 3: Applications -> Jobs
add_heading('Connection 3: My Applications table links to my own Jobs table (M2 -> M3)', level=2)
add_para('How they connect:', bold=True)
add_para(
    'Every application has a `job_id` foreign key pointing to the `jobs` table. '
    'This is how M3 (applications) sits on top of M2 (jobs). You cannot apply '
    'to a job that does not exist or is not active.'
)
add_para('Where you can see this connection in my code:', bold=True)
add_bullet('In ApplicationController.php create() - I check the job exists and is_active = 1 before inserting')
add_bullet('In ApplicationController.php index() - I JOIN with jobs to show job_title, job_type, and deadline')
add_bullet('In BrowseJobs.vue - the Apply button sends the job_id with the application')
add_para(
    'In simple words: My application system is built on top of my job system. '
    'You apply for a job, so every application must reference a valid, active job.'
)

add_para('')

# Connection 4: Applications -> Users
add_heading('Connection 4: My Applications table links to Areeb\'s Users table (M1)', level=2)
add_para('How they connect:', bold=True)
add_para(
    'Every application has a `user_id` foreign key pointing to the `users` table. '
    'This records which student applied. The user_id comes from the JWT token. '
    'When listing applications, students only see their own (filtered by user_id). '
    'When withdrawing, the system checks that the application belongs to the '
    'logged-in student (ownership check).'
)
add_para('Where you can see this connection in my code:', bold=True)
add_bullet('In ApplicationController.php index() - if role is student, filter by user_id from JWT')
add_bullet('In ApplicationController.php create() - user_id comes from JWT payload')
add_bullet('In ApplicationController.php delete() - ownership check: user_id must match')
add_bullet('In ApplicationController.php index() - JOIN with users to show applicant_name and email')
add_para(
    'In simple words: Applications are tied to the logged-in student. '
    'Students can only see and withdraw their own applications. '
    'The user ID comes from Areeb\'s login token, so it cannot be faked.'
)

add_para('')

# Connection 5: Applications -> Companies (through Jobs)
add_heading('Connection 5: My Applications link to Mariam\'s Companies (through Jobs)', level=2)
add_para('How they connect:', bold=True)
add_para(
    'When listing applications, I JOIN through the jobs table to the companies table '
    'to show the company name. This is a two-hop connection: '
    'applications -> jobs (my tables) -> companies (Mariam\'s table).'
)
add_para('The SQL query in ApplicationController.php index():', bold=True)
add_code('''SELECT a.application_id, a.status, a.applied_at,
       j.title AS job_title,           -- from my jobs table
       c.name AS company_name,         -- from Mariam's companies table
       u.full_name AS applicant_name   -- from Areeb's users table
FROM applications a
JOIN jobs j ON j.job_id = a.job_id
JOIN companies c ON c.company_id = j.company_id
JOIN users u ON u.user_id = a.user_id''')
add_para(
    'In simple words: When showing applications, I pull data from 4 tables at once - '
    'my applications table, my jobs table, Mariam\'s companies table, and Areeb\'s users table. '
    'This gives the user all the info they need in one request.'
)

add_para('')

# Connection 6: Admin Analytics reads my tables
add_heading('Connection 6: Mariam\'s Admin Analytics (M5) reads my tables', level=2)
add_para('How they connect:', bold=True)
add_para(
    'Mariam built the admin dashboard with charts. Her AdminController::analytics() method '
    'counts records from my `jobs` and `applications` tables to show statistics like '
    'total jobs, active jobs, total applications, and applications grouped by status. '
    'I did not write this code, but my tables provide the data.'
)
add_para('What Mariam\'s code does with my tables:', bold=True)
add_code('''// From AdminController.php (Mariam's code)
SELECT COUNT(*) FROM jobs;                          -> counts my jobs
SELECT COUNT(*) FROM jobs WHERE is_active = 1;      -> counts my active jobs
SELECT COUNT(*) FROM applications;                   -> counts my applications
SELECT COUNT(*) FROM applications
  WHERE status = 'pending';                          -> groups my applications by status''')
add_para(
    'In simple words: Mariam\'s dashboard reads from my tables to show charts and stats. '
    'My data feeds her analytics. This is a read-only connection - she does not modify my data.'
)

add_para('')

# Connection 7: Shared infrastructure
add_heading('Connection 7: Shared infrastructure I use (built by Areeb)', level=2)
add_para(
    'My modules rely on several shared files that Areeb built. I did not write these, '
    'but my code would not work without them:'
)
add_table(
    ['File', 'Who Built It', 'What It Does for My Modules'],
    [
        ['backend/src/Config/Database.php', 'Areeb (M1)',
         'Creates the PDO database connection. My controllers call Database::getConnection() to talk to MySQL.'],
        ['backend/src/Middleware/JwtMiddleware.php', 'Areeb (M1)',
         'Checks JWT tokens and roles. My routes use this to ensure only admins can create/edit/delete jobs, '
         'and only students can apply.'],
        ['backend/src/Middleware/CorsMiddleware.php', 'Areeb (M1)',
         'Allows the frontend to make requests to the backend. Without this, the browser would block all my API calls.'],
        ['backend/src/Support/Json.php', 'Areeb (M1)',
         'Helper that formats JSON responses. Every one of my endpoints uses Json::write() to return data.'],
        ['frontend/src/services/api.js', 'Areeb (M1)',
         'Axios instance with JWT interceptor. My Vue views use this to make API calls. '
         'It automatically attaches the JWT token to every request.'],
        ['frontend/src/stores/auth.js', 'Areeb (M1)',
         'Pinia store for login state. My views check this to know if the user is logged in and what role they have.'],
        ['frontend/src/router/index.js', 'Areeb (M1)',
         'Vue Router with navigation guards. My views are registered here with role-based access control.'],
        ['frontend/src/components/JobCard.vue', 'Areeb (M1)',
         'Reusable job card component. My BrowseJobs.vue uses this to render each job in the grid.'],
        ['frontend/src/components/StatusBadge.vue', 'Areeb (M1)',
         'Colored status pill. My MyApplications.vue and ManageApplications.vue use this to show application status.'],
        ['frontend/src/components/ConfirmDialog.vue', 'Areeb (M6)',
         '"Are you sure?" modal. My ManageJobs.vue and MyApplications.vue use this before delete/withdraw actions.'],
        ['frontend/src/composables/useToast.js', 'Areeb (M6)',
         'Toast notification helper. My views use toast.success() and toast.error() to show pop-up messages.'],
    ]
)

add_page_break()

# ============================================================================
# 4. DATABASE TABLES
# ============================================================================

add_heading('4. The Database Tables I Designed', level=1)

add_para('I designed two tables in the database. Here they are in detail:')

add_heading('Table 1: jobs (Module 2)', level=2)
add_table(
    ['Column', 'Type', 'What It Stores', 'Notes'],
    [
        ['job_id', 'INT AUTO_INCREMENT', 'Unique ID for each job', 'Primary Key'],
        ['company_id', 'INT', 'Which company posted this job', 'Foreign Key -> companies.company_id (Mariam\'s table)'],
        ['posted_by', 'INT', 'Which admin created this job', 'Foreign Key -> users.user_id (Areeb\'s table)'],
        ['title', 'VARCHAR(150)', 'Job title (e.g., "Software Engineer")', 'Max 150 characters'],
        ['type', 'ENUM', 'Job type', 'Only allows: internship, fulltime, parttime'],
        ['description', 'TEXT', 'Full job description', 'Required, no length limit'],
        ['requirements', 'TEXT', 'What the candidate needs', 'Optional, can be NULL'],
        ['deadline', 'DATE', 'Application deadline', 'Must be a future date (validated in backend)'],
        ['is_active', 'TINYINT(1)', 'Is this job visible to students?', '1 = visible, 0 = hidden (soft delete)'],
        ['created_at', 'TIMESTAMP', 'When the job was created', 'Automatically set to current time'],
    ]
)

add_para('')
add_heading('Table 2: applications (Module 3)', level=2)
add_table(
    ['Column', 'Type', 'What It Stores', 'Notes'],
    [
        ['application_id', 'INT AUTO_INCREMENT', 'Unique ID for each application', 'Primary Key'],
        ['job_id', 'INT', 'Which job was applied for', 'Foreign Key -> jobs.job_id (my table)'],
        ['user_id', 'INT', 'Which student applied', 'Foreign Key -> users.user_id (Areeb\'s table)'],
        ['cover_letter', 'TEXT', 'Optional message from student', 'Max 5000 characters, can be NULL'],
        ['status', 'ENUM', 'Current status of the application', 'Only allows: pending, reviewed, accepted, rejected'],
        ['applied_at', 'TIMESTAMP', 'When the student applied', 'Automatically set to current time'],
        ['updated_at', 'TIMESTAMP', 'When the status was last changed', 'Auto-updates on modification'],
    ]
)
add_para('')
add_para('Special constraint: UNIQUE KEY (job_id, user_id)', bold=True)
add_para(
    'This means the combination of job_id and user_id must be unique. '
    'A student can apply to many different jobs, and a job can receive many applications, '
    'but the SAME student cannot apply to the SAME job twice. This is enforced at the '
    'database level, so even if someone bypasses my code, the database will reject the duplicate.'
)

add_para('')
add_para('How my tables relate to each other and to teammates\' tables:', bold=True)
add_code('''
Areeb's users table (M1)
    |
    |-- posted_by -----> My jobs table (M2)
    |                       |
    |                       |-- company_id --> Mariam's companies table (M4)
    |
    |-- user_id --------> My applications table (M3)
                            |
                            |-- job_id ------> My jobs table (M2)

Mariam's AdminController (M5) reads from:
    |-- My jobs table (counts total and active jobs)
    |-- My applications table (counts by status for the dashboard chart)
''')

add_page_break()

# ============================================================================
# 5. API ENDPOINTS
# ============================================================================

add_heading('5. My API Endpoints (What URLs My Code Handles)', level=1)

add_para('My code handles 9 API endpoints. Here is each one:')

add_table(
    ['Method', 'URL', 'What It Does', 'Who Can Access', 'My File'],
    [
        ['GET', '/api/jobs', 'List all active jobs (with search and filter)', 'Anyone (public)', 'JobController::index'],
        ['GET', '/api/jobs/{id}', 'Get details of one job', 'Anyone (public)', 'JobController::show'],
        ['POST', '/api/jobs', 'Create a new job', 'Admin / Super Admin', 'JobController::create'],
        ['PUT', '/api/jobs/{id}', 'Update an existing job', 'Admin / Super Admin', 'JobController::update'],
        ['DELETE', '/api/jobs/{id}', 'Deactivate a job (soft delete)', 'Admin / Super Admin', 'JobController::delete'],
        ['GET', '/api/applications', 'List applications (students see own, admins see all)', 'Any logged-in user', 'ApplicationController::index'],
        ['POST', '/api/applications', 'Apply for a job', 'Student only', 'ApplicationController::create'],
        ['PUT', '/api/applications/{id}/status', 'Change application status', 'Admin / Super Admin', 'ApplicationController::updateStatus'],
        ['DELETE', '/api/applications/{id}', 'Withdraw an application', 'Student (own + pending only)', 'ApplicationController::delete'],
    ]
)

add_para('')
add_para('How routes are registered in index.php (with Areeb\'s JWT middleware):', bold=True)
add_code('''// M2: Jobs - GET is public, POST/PUT/DELETE need admin role
$app->get('/api/jobs',           [JobController::class, 'index']);
$app->get('/api/jobs/{id}',      [JobController::class, 'show']);
$app->post('/api/jobs',          [JobController::class, 'create'])
    ->add(new JwtMiddleware(['admin', 'superadmin']));
$app->put('/api/jobs/{id}',      [JobController::class, 'update'])
    ->add(new JwtMiddleware(['admin', 'superadmin']));
$app->delete('/api/jobs/{id}',   [JobController::class, 'delete'])
    ->add(new JwtMiddleware(['admin', 'superadmin']));

// M3: Applications - all need login, different roles for different actions
$app->get('/api/applications',   [ApplicationController::class, 'index'])
    ->add(new JwtMiddleware());
$app->post('/api/applications',  [ApplicationController::class, 'create'])
    ->add(new JwtMiddleware('student'));
$app->put('/api/applications/{id}/status', [ApplicationController::class, 'updateStatus'])
    ->add(new JwtMiddleware(['admin', 'superadmin']));
$app->delete('/api/applications/{id}', [ApplicationController::class, 'delete'])
    ->add(new JwtMiddleware('student'));''')

add_page_break()

# ============================================================================
# 6. M2 BACKEND CODE EXPLAINED
# ============================================================================

add_heading('6. Module 2: Job Management - Backend Code Explained', level=1)
add_para('File: backend/src/Modules/Jobs/JobController.php')
add_para('This file has 6 methods. Here is each one explained in simple terms:')

# 6a
add_heading('6a. Listing Jobs (index method) - GET /api/jobs', level=2)
add_para('What it does:', bold=True)
add_para(
    'Returns all active jobs from the database. Supports optional search (by title or company name) '
    'and optional filter (by job type). Results are sorted by newest first.'
)
add_para('How it works step by step:', bold=True)
add_number('Read query parameters from the URL: search, type, company_id')
add_number('Build a SQL query that starts with "SELECT ... FROM jobs JOIN companies ..."')
add_number('If search is provided, add a WHERE clause with LIKE for title and company name')
add_number('If type is provided, add a WHERE clause filtering by type')
add_number('Always filter by is_active = 1 (only show active jobs)')
add_number('Execute the query using PDO prepared statements (safe from SQL injection)')
add_number('Return the results as JSON with status 200')
add_para('Key code:', bold=True)
add_code('''$sql = "SELECT j.job_id, j.title, j.type, j.description, j.deadline,
               c.name AS company_name, c.location AS company_location
        FROM jobs j
        JOIN companies c ON c.company_id = j.company_id
        WHERE j.is_active = 1";

if ($search !== '') {
    $sql .= ' AND (j.title LIKE :qt OR c.name LIKE :qc)';
    $args[':qt'] = '%' . $search . '%';
    $args[':qc'] = '%' . $search . '%';
}
if (in_array($type, ['internship', 'fulltime', 'parttime'], true)) {
    $sql .= ' AND j.type = :type';
    $args[':type'] = $type;
}

$stmt = Database::getConnection()->prepare($sql);
$stmt->execute($args);
$rows = $stmt->fetchAll(PDO::FETCH_ASSOC);

return Json::write($response, 200, ['success' => true, 'data' => $rows]);''')
add_para('In simple words: It builds a flexible SQL query that can search and filter, '
         'then safely runs it with PDO. The JOIN with companies means each job comes with '
         'its company name - no need for a second request.')

# 6b
add_heading('6b. Viewing a Single Job (show method) - GET /api/jobs/{id}', level=2)
add_para('What it does:', bold=True)
add_para('Returns the full details of one job, including company info and who posted it.')
add_para('Key code:', bold=True)
add_code('''$stmt = Database::getConnection()->prepare(
    "SELECT j.*, c.name AS company_name, c.industry AS company_industry,
            c.location AS company_location, u.full_name AS posted_by_name
     FROM jobs j
     JOIN companies c ON c.company_id = j.company_id
     JOIN users u ON u.user_id = j.posted_by
     WHERE j.job_id = :id LIMIT 1"
);
$stmt->execute([':id' => $id]);''')
add_para('In simple words: It joins 3 tables (jobs, companies, users) to get everything '
         'needed for the job detail page in one query. If the job does not exist, returns 404.')

# 6c
add_heading('6c. Creating a Job (create method) - POST /api/jobs', level=2)
add_para('What it does:', bold=True)
add_para('Creates a new job listing. Only admins can access this (enforced by JWT middleware).')
add_para('How it works step by step:', bold=True)
add_number('Read the JSON body from the request (title, company_id, type, description, requirements, deadline)')
add_number('Get the user_id from the JWT token (this becomes posted_by)')
add_number('Validate all fields using the validate() method')
add_number('If validation fails, return 400 with error messages')
add_number('Insert the job using PDO prepared statement with named placeholders')
add_number('Get the new job ID using lastInsertId()')
add_number('Fetch the new job with company name and return it as JSON with status 201')
add_para('Key code:', bold=True)
add_code('''$payload = (array) $request->getAttribute('jwt_payload');

$stmt = $pdo->prepare(
    'INSERT INTO jobs (company_id, posted_by, title, type,
         description, requirements, deadline, is_active)
     VALUES (:cid, :uid, :title, :type,
         :descr, :req, :deadline, 1)'
);
$stmt->execute([
    ':cid'      => (int) $body['company_id'],
    ':uid'      => (int) $payload['user_id'],  // from JWT, not from user input
    ':title'    => $body['title'],
    ':type'     => $body['type'],
    ':descr'    => $body['description'],
    ':req'      => $body['requirements'] ?? null,
    ':deadline' => $body['deadline'],
]);''')
add_para('In simple words: The admin fills a form, the backend validates it, then inserts it '
         'into the database. The posted_by field comes from the login token, so the system '
         'always knows who created the job.')

# 6d
add_heading('6d. Updating a Job (update method) - PUT /api/jobs/{id}', level=2)
add_para('What it does:', bold=True)
add_para('Updates an existing job. Only admins can access this.')
add_para('How it works:', bold=True)
add_number('Validate the job ID from the URL')
add_number('Read and validate the JSON body')
add_number('Check if the job exists (SELECT by ID) - return 404 if not found')
add_number('Update the job using PDO prepared statement')
add_number('Fetch the updated job with company name and return it as JSON')

# 6e
add_heading('6e. Deleting a Job (delete method) - Soft Delete', level=2)
add_para('What it does:', bold=True)
add_para(
    'This is NOT a real DELETE. Instead of removing the job from the database, '
    'it sets is_active to 0. This is called a "soft delete".'
)
add_para('Key code:', bold=True)
add_code('''// Check job exists first
$stmt = $pdo->prepare('SELECT job_id FROM jobs WHERE job_id = :id');
$stmt->execute([':id' => $id]);
if (!$stmt->fetch()) {
    return Json::write($response, 404, [...]);
}

// Soft delete: set is_active = 0 (NOT a real DELETE)
$stmt = $pdo->prepare('UPDATE jobs SET is_active = 0 WHERE job_id = :id');
$stmt->execute([':id' => $id]);''')
add_para('Why soft delete instead of hard delete?', bold=True)
add_bullet('Existing applications still reference the job. If we hard-deleted it, those references would break.')
add_bullet('The job data is preserved for audit and reporting purposes.')
add_bullet('Students just do not see it anymore because the index query filters WHERE is_active = 1.')
add_bullet('Admins can see it was deactivated (is_active = 0 in the database).')

# 6f
add_heading('6f. Input Validation (validate method)', level=2)
add_para('What it does:', bold=True)
add_para('Checks all fields before creating or updating a job. Returns an array of error messages.')
add_para('What it validates:', bold=True)
add_bullet('company_id: must be a valid number')
add_bullet('title: required, max 150 characters')
add_bullet('type: must be one of internship, fulltime, or parttime')
add_bullet('description: required')
add_bullet('deadline: required, must be in YYYY-MM-DD format, must be a future date')
add_para('Key code:', bold=True)
add_code('''$deadline = (string) ($body['deadline'] ?? '');
if ($deadline === '') {
    $errors['deadline'] = 'Deadline is required.';
} else {
    $dt = DateTime::createFromFormat('Y-m-d', $deadline);
    if (!$dt || $dt->format('Y-m-d') !== $deadline) {
        $errors['deadline'] = 'Deadline must be in YYYY-MM-DD format.';
    } elseif ($dt < new DateTime('today')) {
        $errors['deadline'] = 'Deadline must be a future date.';
    }
}''')

add_page_break()

# ============================================================================
# 7. M3 BACKEND CODE EXPLAINED
# ============================================================================

add_heading('7. Module 3: Application Tracking - Backend Code Explained', level=1)
add_para('File: backend/src/Modules/Applications/ApplicationController.php')
add_para('This file has 4 methods. Here is each one explained in simple terms:')

# 7a
add_heading('7a. Listing Applications (index method) - GET /api/applications', level=2)
add_para('What it does:', bold=True)
add_para(
    'Returns applications. This endpoint is "role-aware" - it behaves differently '
    'depending on who is asking:'
)
add_bullet('If a student calls it: returns only THEIR applications (filtered by user_id)')
add_bullet('If an admin calls it: returns ALL applications in the system')
add_para('How it works:', bold=True)
add_number('Get user_id and role from the JWT token')
add_number('Build a SQL query that JOINs applications, jobs, companies, and users')
add_number('If role is student, add WHERE a.user_id = :uid to filter by their own ID')
add_number('If a status filter is provided (e.g., ?status=pending), add a WHERE clause')
add_number('Execute and return as JSON')
add_para('Key code:', bold=True)
add_code('''$payload = (array) $request->getAttribute('jwt_payload');
$userId  = (int) ($payload['user_id'] ?? 0);
$role    = (string) ($payload['role'] ?? '');

$sql = "SELECT a.application_id, a.status, a.applied_at,
               j.title AS job_title, c.name AS company_name,
               u.full_name AS applicant_name, u.email AS applicant_email
        FROM applications a
        JOIN jobs j ON j.job_id = a.job_id
        JOIN companies c ON c.company_id = j.company_id
        JOIN users u ON u.user_id = a.user_id
        WHERE 1=1";

if ($role === 'student') {
    $sql .= ' AND a.user_id = :uid';
    $args[':uid'] = $userId;
}''')
add_para('In simple words: The same endpoint serves both students and admins. '
         'Students see only their own applications. Admins see everything. '
         'The role comes from the JWT token, so it cannot be faked.')

# 7b
add_heading('7b. Applying for a Job (create method) - POST /api/applications', level=2)
add_para('What it does:', bold=True)
add_para('Lets a student apply for a job. Only students can access this.')
add_para('How it works step by step:', bold=True)
add_number('Get user_id from JWT token')
add_number('Read job_id and cover_letter from the request body')
add_number('Validate: job_id must be a positive number, cover_letter max 5000 chars')
add_number('Check 1: Does the job exist? (SELECT from jobs table) - return 404 if not')
add_number('Check 2: Is the job still active? (is_active = 1) - return 400 if deactivated')
add_number('Check 3: Has this student already applied? (SELECT from applications) - return 409 if duplicate')
add_number('Insert the application with status "pending" using PDO prepared statement')
add_number('Return 201 Created with the new application ID')
add_para('Key code (the 3 checks):', bold=True)
add_code('''// Check 1: job exists
$stmt = $pdo->prepare('SELECT job_id, is_active FROM jobs WHERE job_id = :id');
$stmt->execute([':id' => $jobId]);
$job = $stmt->fetch(PDO::FETCH_ASSOC);
if (!$job) return Json::write($response, 404, [...]);

// Check 2: job is active
if ((int) $job['is_active'] !== 1)
    return Json::write($response, 400, [...]);

// Check 3: not already applied
$stmt = $pdo->prepare(
    'SELECT application_id FROM applications
     WHERE job_id = :jid AND user_id = :uid LIMIT 1'
);
$stmt->execute([':jid' => $jobId, ':uid' => $userId]);
if ($stmt->fetch())
    return Json::write($response, 409, [
        'success' => false,
        'message' => 'You have already applied for this job.',
    ]);

// All checks passed - insert the application
$stmt = $pdo->prepare(
    'INSERT INTO applications (job_id, user_id, cover_letter, status)
     VALUES (:jid, :uid, :cl, "pending")'
);
$stmt->execute([...]);''')
add_para('In simple words: Three safety checks before inserting. The job must exist, '
         'must be active, and the student must not have already applied. '
         'Only then is the application saved with status "pending".')

# 7c
add_heading('7c. Changing Application Status (updateStatus method) - PUT', level=2)
add_para('What it does:', bold=True)
add_para('Lets an admin change the status of an application. Only admins can access this.')
add_para('How it works:', bold=True)
add_number('Validate the application ID from the URL')
add_number('Validate the status value (must be pending, reviewed, accepted, or rejected)')
add_number('Check if the application exists - return 404 if not')
add_number('Update the status using PDO prepared statement')
add_number('Return 200 with the new status')
add_para('In simple words: The admin picks a new status from a dropdown, '
         'the backend validates it is one of the four allowed values, then updates the database.')

# 7d
add_heading('7d. Withdrawing an Application (delete method) - DELETE', level=2)
add_para('What it does:', bold=True)
add_para('Lets a student withdraw (delete) their application. Only students can access this, '
         'and only for their own pending applications.')
add_para('Three security checks before deleting:', bold=True)
add_number('Check 1: Does the application exist? - return 404 if not')
add_number('Check 2: Does it belong to this student? - return 403 Forbidden if not')
add_number('Check 3: Is it still "pending"? - return 400 if already reviewed/accepted/rejected')
add_para('Key code:', bold=True)
add_code('''$stmt = $pdo->prepare(
    'SELECT user_id, status FROM applications
     WHERE application_id = :id LIMIT 1'
);
$stmt->execute([':id' => $id]);
$row = $stmt->fetch(PDO::FETCH_ASSOC);

if (!$row) return Json::write($response, 404, [...]);

// Ownership check - not your application
if ((int) $row['user_id'] !== $userId) {
    return Json::write($response, 403, [
        'success' => false,
        'message' => 'You can only withdraw your own applications.',
    ]);
}

// Status check - only pending can be withdrawn
if ($row['status'] !== 'pending') {
    return Json::write($response, 400, [
        'success' => false,
        'message' => 'Only pending applications can be withdrawn.',
    ]);
}

// All checks passed - hard delete
$stmt = $pdo->prepare('DELETE FROM applications WHERE application_id = :id');
$stmt->execute([':id' => $id]);''')
add_para('In simple words: A student can only withdraw their own application, and only if '
         'the admin has not reviewed it yet. Once the status changes from "pending", '
         'the student cannot withdraw it anymore. This prevents students from hiding '
         'rejected applications.')

add_page_break()

# ============================================================================
# 8. FRONTEND CODE
# ============================================================================

add_heading('8. My Frontend Code (Vue.js Views)', level=1)

add_heading('8a. BrowseJobs.vue (M2 - Student)', level=2)
add_para('What it does:', bold=True)
add_bullet('Shows a grid of job cards with search and filter controls')
add_bullet('On page load, fetches both jobs (GET /api/jobs) and applications (GET /api/applications) in parallel')
add_bullet('Uses the applications data to disable the Apply button for jobs the student already applied to')
add_bullet('Search has 400ms debounce (waits until user stops typing before filtering)')
add_bullet('Filtering is done with a Vue computed property (no page reload needed)')
add_bullet('Clicking Apply opens a modal with a cover letter textarea (max 5000 chars)')
add_bullet('On submit, sends POST /api/applications and shows a success toast')
add_para('Key frontend pattern - parallel loading:', bold=True)
add_code('''const [jobsRes, appsRes] = await Promise.all([
  api.get('/jobs'),
  api.get('/applications').catch(() => ({ data: { data: [] } }))
])
jobs.value = jobsRes.data?.data ?? []
appliedJobIds.value = new Set(
  (appsRes.data?.data ?? []).map(a => Number(a.job_id))
)''')
add_para('In simple words: The page loads jobs and the student\'s existing applications at the '
         'same time (not one after the other). This makes the page load faster. The applications '
         'are used to grey out the Apply button for jobs the student already applied to.')

add_heading('8b. JobDetail.vue (M2 - Student)', level=2)
add_para('What it does:', bold=True)
add_bullet('Shows full details of one job (description, requirements, deadline, company info)')
add_bullet('Calls GET /api/jobs/{id} when the page loads')
add_bullet('Has an Apply button that opens the same apply modal as BrowseJobs')

add_heading('8c. ManageJobs.vue (M2 - Admin)', level=2)
add_para('What it does:', bold=True)
add_bullet('Shows a table of all jobs with Edit and Delete buttons')
add_bullet('Add Job button opens a modal form with: title, company dropdown, type, deadline, description, requirements')
add_bullet('The company dropdown is populated by calling GET /api/companies (Mariam\'s API)')
add_bullet('Frontend validation: required fields, title max 150 chars, deadline must be in the future')
add_bullet('Edit fills the form with existing job data and sends PUT /api/jobs/{id}')
add_bullet('Delete shows a confirm dialog, then sends DELETE /api/jobs/{id} (soft delete)')
add_bullet('After delete, the job stays in the table but is_active changes to 0 (shown as inactive)')
add_para('Key frontend pattern - form validation:', bold=True)
add_code('''function validate() {
  if (!form.company_id) errors.company_id = 'Company is required.'
  if (!form.title.trim()) errors.title = 'Title is required.'
  else if (form.title.length > 150) errors.title = 'Title must be 150 chars or fewer.'
  if (!form.description.trim()) errors.description = 'Description is required.'
  if (!form.deadline) errors.deadline = 'Deadline is required.'
  else if (form.deadline < todayStr.value) errors.deadline = 'Deadline must be in the future.'
  return Object.values(errors).every(v => v === '')
}''')

add_heading('8d. MyApplications.vue (M3 - Student)', level=2)
add_para('What it does:', bold=True)
add_bullet('Shows a table of the student\'s applications with status badges')
add_bullet('Tabs at the top filter by status: All, Pending, Reviewed, Accepted, Rejected')
add_bullet('Each tab shows a count of applications in that status')
add_bullet('Withdraw button only appears for pending applications')
add_bullet('Clicking Withdraw shows a confirm dialog, then sends DELETE /api/applications/{id}')
add_bullet('After withdraw, the application is removed from the table instantly')

add_heading('8e. ManageApplications.vue (M3 - Admin)', level=2)
add_para('What it does:', bold=True)
add_bullet('Shows a table of ALL applications in the system (from all students)')
add_bullet('Shows applicant name, email, job title, company, applied date, and current status')
add_bullet('A dropdown on each row lets the admin change the status')
add_bullet('Uses "optimistic update" - the UI updates instantly, then sends the API request')
add_bullet('If the API request fails, the UI rolls back to the previous status and shows an error toast')
add_para('Key frontend pattern - optimistic update:', bold=True)
add_code('''async function changeStatus(app, newStatus) {
  const previous = app.status
  // 1. Update UI immediately
  const idx = applications.value.findIndex(a => a.application_id === app.application_id)
  if (idx !== -1) applications.value[idx] = { ...app, status: newStatus }

  try {
    // 2. Send API request
    await api.put(`/applications/${app.application_id}/status`, { status: newStatus })
    toast.success(`Status set to ${newStatus}.`)
  } catch (err) {
    // 3. Rollback on error
    if (idx !== -1) applications.value[idx] = { ...app, status: previous }
    toast.error('Could not update status.')
  }
}''')
add_para('In simple words: When the admin changes a status, the dropdown updates instantly '
         '(no loading spinner). The API request happens in the background. If it fails, '
         'the dropdown goes back to the old value and shows an error. This makes the app feel fast.')

add_page_break()

# ============================================================================
# 9. SHARED CODE
# ============================================================================

add_heading('9. Shared Code I Use (Built by Teammates)', level=1)

add_para(
    'My modules use shared infrastructure that Areeb built. I did not write these files, '
    'but my code depends on them. If the lecturer asks "did you build everything yourself?", '
    'the honest answer is: I built the job and application logic, and I used shared tools '
    'that Areeb set up for the whole team.'
)

add_para('')
add_para('What I built vs what I use:', bold=True)
add_table(
    ['What', 'Who Built It', 'How My Code Uses It'],
    [
        ['JobController.php', 'ME', 'My M2 backend logic'],
        ['ApplicationController.php', 'ME', 'My M3 backend logic'],
        ['BrowseJobs.vue, JobDetail.vue, ManageJobs.vue', 'ME', 'My M2 frontend'],
        ['MyApplications.vue, ManageApplications.vue', 'ME', 'My M3 frontend'],
        ['jobs and applications table design', 'ME', 'My database schema'],
        ['Database.php (PDO connection)', 'Areeb', 'I call Database::getConnection() in every method'],
        ['JwtMiddleware.php', 'Areeb', 'I attach it to my routes to enforce role-based access'],
        ['Json.php', 'Areeb', 'I call Json::write() to return all my responses'],
        ['api.js (Axios)', 'Areeb', 'My Vue views use this to make API calls with JWT attached'],
        ['auth.js (Pinia store)', 'Areeb', 'My views check login state from this store'],
        ['router/index.js', 'Areeb', 'My views are registered here with role guards'],
        ['JobCard.vue, StatusBadge.vue', 'Areeb', 'Reusable UI components I use in my views'],
        ['ConfirmDialog.vue, useToast.js', 'Areeb', 'I use these for delete confirms and notifications'],
    ]
)

add_page_break()

# ============================================================================
# 10. SECURITY
# ============================================================================

add_heading('10. Security in My Modules', level=1)

add_heading('Backend Security (in my PHP controllers)', level=2)
add_table(
    ['Security Feature', 'Where in My Code', 'What It Prevents'],
    [
        ['PDO Prepared Statements',
         'Every query in JobController and ApplicationController uses :placeholder syntax',
         'SQL injection - user input is treated as data, not as SQL commands'],
        ['JWT Role Gating',
         'Routes in index.php use JwtMiddleware([\'admin\',\'superadmin\']) or JwtMiddleware(\'student\')',
         'Unauthorized access - only the right roles can access each endpoint'],
        ['Ownership Check',
         'ApplicationController::delete() checks user_id matches the JWT user_id',
         'A student withdrawing another student\'s application (returns 403 Forbidden)'],
        ['Status Check',
         'ApplicationController::delete() checks status is still "pending"',
         'Withdrawing an application that has already been reviewed by an admin'],
        ['Duplicate Prevention',
         'ApplicationController::create() checks if already applied + database UNIQUE constraint',
         'A student applying to the same job twice (returns 409 Conflict)'],
        ['Soft Delete',
         'JobController::delete() sets is_active = 0 instead of DELETE',
         'Data loss - job data is preserved for applications and audit'],
        ['Input Validation',
         'validate() method in JobController checks all fields',
         'Invalid data entering the database (empty titles, past deadlines, invalid types)'],
        ['Job Active Check',
         'ApplicationController::create() checks is_active = 1 before accepting application',
         'Students applying to deactivated jobs'],
    ]
)

add_para('')
add_heading('Frontend Security (in my Vue views)', level=2)
add_table(
    ['Security Feature', 'Where in My Code', 'What It Does'],
    [
        ['Route Guards', 'router/index.js (Areeb\'s file, my routes)',
         'Blocks students from visiting /admin/* pages and blocks admins from /student/* pages'],
        ['Client-side Validation', 'ManageJobs.vue validate() function',
         'Shows error messages before sending the request (better user experience)'],
        ['Confirm Dialogs', 'ManageJobs.vue and MyApplications.vue',
         'Shows "Are you sure?" before delete or withdraw actions'],
        ['Apply Button Disabled', 'BrowseJobs.vue checks appliedJobIds Set',
         'Greys out the Apply button for jobs the student already applied to'],
        ['Toast Notifications', 'All my views use useToast()',
         'Shows success/error messages so the user knows what happened'],
    ]
)

add_page_break()

# ============================================================================
# 11. DATABASE DEMO COMMANDS
# ============================================================================

add_heading('11. Database Demo Commands (If DR Asks to See the Database)', level=1)

add_para('If the lecturer asks to see the database, run these commands in the terminal:')

add_code('''-- 1. Open MySQL and select the database
mysql -u root
USE careerbridge;

-- 2. Show all tables
SHOW TABLES;

-- 3. See the structure of my tables (M2 + M3)
DESCRIBE jobs;
DESCRIBE applications;

-- 4. See the data inside my tables
SELECT * FROM jobs;
SELECT * FROM applications;

-- 5. Show how tables connect (jobs + companies)
SELECT j.job_id, j.title, j.type, c.name AS company,
       j.deadline, j.is_active
FROM jobs j
JOIN companies c ON c.company_id = j.company_id;

-- 6. Show applications with student + job info
SELECT a.application_id, u.full_name AS student,
       j.title AS job, a.status, a.applied_at
FROM applications a
JOIN users u ON u.user_id = a.user_id
JOIN jobs j ON j.job_id = a.job_id;

-- 7. Count applications by status (powers the admin chart)
SELECT status, COUNT(*) AS count
FROM applications GROUP BY status;''')

add_para('')
add_para('Quick one-liner if DR just wants a fast look:', bold=True)
add_code('''mysql -u root -e "USE careerbridge; SHOW TABLES; SELECT * FROM jobs LIMIT 5; SELECT * FROM applications LIMIT 5;"''')

add_page_break()

# ============================================================================
# 12. QUESTIONS THE LECTURER MIGHT ASK
# ============================================================================

add_heading('12. Questions the Lecturer Might Ask (With Answers)', level=1)

add_para('')
add_para('Q: What is CRUD and where is it in your code?', bold=True)
add_para(
    'CRUD stands for Create, Read, Update, Delete. In my M2 (Jobs): Create is POST /api/jobs '
    '(JobController::create), Read is GET /api/jobs and GET /api/jobs/{id} (index and show), '
    'Update is PUT /api/jobs/{id} (update), Delete is DELETE /api/jobs/{id} (delete - soft delete). '
    'In my M3 (Applications): Create is POST /api/applications, Read is GET /api/applications, '
    'Update is PUT /api/applications/{id}/status, Delete is DELETE /api/applications/{id}.'
)

add_para('')
add_para('Q: How do you prevent SQL injection?', bold=True)
add_para(
    'Every database query in my controllers uses PDO prepared statements with named placeholders '
    '(like :id, :title, :cid). The SQL query and the user data are sent separately to MySQL, '
    'so MySQL knows the data is just data, not SQL commands. Even if someone tries to inject '
    'malicious SQL, it will be treated as a regular string value.'
)

add_para('')
add_para('Q: How does your authentication work?', bold=True)
add_para(
    'I use JWT (JSON Web Token) authentication built by Areeb (M1). When a user logs in, '
    'the backend creates a signed token containing their user_id, email, and role. '
    'My routes attach JwtMiddleware which checks this token before allowing access. '
    'For example, POST /api/jobs requires the token to have role "admin" or "superadmin". '
    'POST /api/applications requires role "student". The token comes from the request '
    'header, so it cannot be faked.'
)

add_para('')
add_para('Q: What is soft delete and why did you use it?', bold=True)
add_para(
    'Soft delete means setting is_active = 0 instead of running DELETE. I used it for jobs '
    'because applications reference jobs via foreign key. If I hard-deleted a job, '
    'the applications for that job would have a broken reference. With soft delete, '
    'the job stays in the database, applications still work, and the job is just hidden '
    'from students (the index query filters WHERE is_active = 1).'
)

add_para('')
add_para('Q: How does a student apply for a job? Walk me through the flow.', bold=True)
add_para(
    '1) Student visits BrowseJobs page. 2) Frontend loads jobs (GET /api/jobs) and their '
    'existing applications (GET /api/applications) in parallel. 3) Jobs the student already '
    'applied to show a disabled Apply button. 4) Student clicks Apply on a new job, '
    'a modal opens with a cover letter field. 5) Student writes a cover letter and clicks Submit. '
    '6) Frontend sends POST /api/applications with job_id and cover_letter. '
    '7) Backend checks: job exists? job active? not already applied? '
    '8) If all checks pass, inserts with status "pending" and returns 201. '
    '9) Frontend shows success toast and disables the Apply button for that job.'
)

add_para('')
add_para('Q: Can a student apply to the same job twice?', bold=True)
add_para(
    'No. This is prevented in two places: (1) In my code - ApplicationController::create() '
    'checks if an application already exists for this job_id and user_id, and returns 409 Conflict. '
    '(2) In the database - there is a UNIQUE constraint on (job_id, user_id) in the applications '
    'table, so even if someone bypasses my code, MySQL will reject the duplicate.'
)

add_para('')
add_para('Q: How does your module connect to other team members\' modules?', bold=True)
add_para(
    'My jobs table links to Mariam\'s companies table (company_id foreign key) and Areeb\'s '
    'users table (posted_by foreign key). My applications table links to my own jobs table '
    '(job_id) and Areeb\'s users table (user_id). When listing applications, I JOIN through '
    '4 tables at once: applications, jobs, companies, and users. Mariam\'s admin analytics '
    'reads from my jobs and applications tables to show charts and statistics.'
)

add_para('')
add_para('Q: What HTTP status codes do you use and why?', bold=True)
add_para(
    '200 OK for successful reads/updates. 201 Created for successful POST (new job or application). '
    '400 Bad Request for validation errors. 401 Unauthorized for missing/invalid JWT token. '
    '403 Forbidden for wrong role or not-owner. 404 Not Found for missing records. '
    '409 Conflict for duplicate applications. 500 Server Error for database problems. '
    'These are standard HTTP codes that tell the frontend exactly what happened.'
)

add_para('')
add_para('Q: What is optimistic update and where did you use it?', bold=True)
add_para(
    'Optimistic update is a UI pattern where you update the screen immediately, then send '
    'the API request in the background. I used it in ManageApplications.vue for changing '
    'application status. When the admin picks a new status from the dropdown, the UI updates '
    'instantly. If the API call fails, the UI rolls back to the old status and shows an error. '
    'This makes the app feel fast and responsive.'
)

add_para('')
add_para('Q: Why do you validate on both frontend and backend?', bold=True)
add_para(
    'Frontend validation gives instant feedback to the user (no waiting for the server). '
    'But frontend validation can be bypassed - someone could send a request directly to '
    'the API using curl or Postman. So the backend validates everything again for security. '
    'Both sides check: required fields, max lengths, valid enum values, and future deadline dates.'
)

# ============================================================================
# SAVE
# ============================================================================

output_path = os.path.join(os.path.dirname(__file__), 'Samin_M2_M3_Explanation.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
