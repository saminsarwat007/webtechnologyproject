#!/usr/bin/env python3
"""
Build a comprehensive DOCX file explaining the CareerBridge codebase
in simple, easy-to-understand terms for a beginner preparing for a presentation.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ---- Style helpers ---------------------------------------------------------

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Bullet 2')
    return p

def add_number(text):
    p = doc.add_paragraph(text, style='List Number')
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # light gray shading
    shading = p.paragraph_format.element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5'
    })
    shading.append(shd)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

def add_page_break():
    doc.add_page_break()

# ============================================================================
# TITLE PAGE
# ============================================================================

title = doc.add_heading('', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CareerBridge\nStudent Internship & Job Application Management System')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Complete Codebase Explanation & Presentation Guide')
run.font.size = Pt(14)
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nSECJ3483 Web Technology - Group Project\nTeam CHAMPION\n')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Team Members:\n')
run.font.size = Pt(11)
run.bold = True

add_para('Mohammad Areeb (A22EC4041) - M1 Auth & Access Control, M6 System Admin')
add_para('Samin Sarwat (A22EC4040) - M2 Job Management, M3 Application Tracking')
add_para('Mariam Hanif (A22EC4034) - M4 Company Management, M5 Reporting & Analytics')
add_para('Monika Zelenkov (A22EC4045) - M7 Forum & Discussion, M8 Mock Interview Scheduler')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nThis document explains every part of the codebase in simple terms.\nRead it before your presentation to understand and answer questions.')
run.font.size = Pt(11)
run.italic = True

add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================

add_heading('Table of Contents', level=1)

contents = [
    '1. What is CareerBridge? (Project Overview)',
    '2. Technology Stack (What Tools We Used)',
    '3. How the Project is Organized (Folder Structure)',
    '4. How the Frontend, Backend, and Database Connect',
    '5. Database Setup - How to Create and Fill the Database',
    '6. Database Tables Explained (Schema)',
    '7. Backend Explanation (PHP / Slim 4 REST API)',
    '    7a. Entry Point (index.php) - How Routes Work',
    '    7b. Database Connection (Database.php)',
    '    7c. JWT Authentication (JwtMiddleware.php)',
    '    7d. CORS (CorsMiddleware.php)',
    '    7e. JSON Response Helper (Json.php)',
    '    7f. Module 1: Auth Controller (Login & Register)',
    '    7g. Module 2: Job Controller (CRUD for Jobs)',
    '    7h. Module 3: Application Controller (Apply & Track)',
    '    7i. Module 4: Company Controller (CRUD for Companies)',
    '    7j. Module 5: Profile Controller (Student Profiles)',
    '    7k. Module 5+6: Admin Controller (Analytics & Users)',
    '    7l. Module 7: Forum Controller (Posts, Comments, Likes)',
    '    7m. Module 8: Interview Controller (Slots & Bookings)',
    '8. Frontend Explanation (Vue.js 3 SPA)',
    '    8a. Main Entry (main.js & App.vue)',
    '    8b. Router (How Pages Navigate)',
    '    8c. Auth Store (Pinia - Login State Management)',
    '    8d. API Service (Axios - How Frontend Talks to Backend)',
    '    8e. Components (Reusable UI Parts)',
    '    8f. Views (Page-by-Page Explanation)',
    '9. JWT Authentication Flow (Step by Step)',
    '10. CRUD Operations Summary (Where is GET, POST, PUT, DELETE?)',
    '11. Input Validation (Frontend and Backend)',
    '12. Error Handling & HTTP Status Codes',
    '13. Security Features',
    '14. How to Run the Project (Step by Step)',
    '15. Demo Login Credentials',
    '16. Questions the Lecturer Might Ask (With Answers)',
    '17. Rubric Mapping (How We Meet Each Criterion)',
]

for item in contents:
    add_para(item, size=11)

add_page_break()

# ============================================================================
# 1. PROJECT OVERVIEW
# ============================================================================

add_heading('1. What is CareerBridge? (Project Overview)', level=1)

add_para(
    'CareerBridge is a web application that connects university students with '
    'internship and job opportunities. Think of it like a simplified LinkedIn '
    'or JobStreet, but specifically designed for university students.'
)

add_para('The system has three types of users:', bold=True)

add_table(
    ['Role', 'What They Can Do'],
    [
        ['Student', 'Browse jobs, apply with a cover letter, track application status, '
                    'participate in forum discussions, book mock interviews, manage their profile'],
        ['Admin', 'Post and manage jobs, review applications and change their status, '
                  'manage companies, manage mock interview slots, evaluate interviews'],
        ['Super Admin', 'Everything an admin can do, PLUS view and manage all user accounts'],
    ]
)

add_para('')
add_para('The project is divided into 8 modules, each owned by one team member:', bold=True)

add_table(
    ['Module', 'Name', 'Owner', 'What It Does'],
    [
        ['M1', 'User Authentication & Access Control', 'Areeb',
         'Login, register, JWT tokens, role-based access control'],
        ['M2', 'Job & Internship Management', 'Samin',
         'Create, read, update, delete job listings (full CRUD)'],
        ['M3', 'Application Tracking System', 'Samin',
         'Students apply for jobs, admins change application status, students withdraw'],
        ['M4', 'Company & External Source Management', 'Mariam',
         'Create, read, update, delete companies (full CRUD)'],
        ['M5', 'Reporting & Analytics', 'Mariam',
         'Admin dashboard with charts, student profiles, statistics'],
        ['M6', 'System Administration & Notifications', 'Areeb',
         'Super admin user directory, toast notifications, confirm dialogs'],
        ['M7', 'Forum & Discussion', 'Monika',
         'Forum posts with tags, comments, likes (toggle), search and filter'],
        ['M8', 'Mock Interview & Technical Prep Scheduler', 'Monika',
         'Admins create interview slots, students book them, admins give scores and feedback'],
    ]
)

add_page_break()

# ============================================================================
# 2. TECH STACK
# ============================================================================

add_heading('2. Technology Stack (What Tools We Used)', level=1)

add_para('We used different technologies for different parts of the project. Here is what each one does in simple terms:')

add_table(
    ['Layer', 'Technology', 'Simple Explanation'],
    [
        ['Frontend', 'Vue.js 3', 'A JavaScript framework for building user interfaces. '
                     'It lets us create reusable components (like buttons, cards, forms) and '
                     'show them on web pages. Version 3 is the latest major version.'],
        ['Frontend Build', 'Vite', 'A tool that builds our Vue.js code quickly. '
                            'It replaces older tools like webpack. It also provides a '
                            'development server that auto-refreshes when we save files.'],
        ['State Management', 'Pinia', 'A library that stores data we need across multiple pages '
                                '(like who is logged in). Think of it as a shared memory '
                                'for the whole app.'],
        ['Routing', 'Vue Router', 'Lets us navigate between pages without reloading the browser. '
                          'This is what makes it a "Single Page Application" (SPA).'],
        ['Styling', 'Tailwind CSS', 'A CSS framework with pre-made utility classes. '
                            'Instead of writing custom CSS, we use classes like "bg-blue-500" '
                            'or "text-center" directly in our HTML.'],
        ['Charts', 'Chart.js', 'A library for drawing charts. We use it on the admin dashboard '
                       'to show a bar chart of application statuses.'],
        ['HTTP Client', 'Axios', 'A library for making API requests (HTTP calls) from the '
                         'frontend to the backend. It is like fetch() but with more features.'],
        ['Backend', 'PHP (Slim 4)', 'PHP is the programming language for our backend. '
                       'Slim 4 is a lightweight framework that helps us organize our API routes '
                       'and handle HTTP requests and responses.'],
        ['Auth', 'JWT (firebase/php-jwt)', 'JSON Web Tokens are used for authentication. '
                      'When a user logs in, the backend creates a signed token. The frontend '
                      'sends this token with every request to prove who the user is.'],
        ['Database', 'MySQL', 'A relational database that stores all our data in tables '
                      '(users, jobs, applications, etc.). We use PDO (PHP Data Objects) '
                      'to safely talk to the database.'],
        ['Env Config', 'vlucas/phpdotenv', 'A library that loads environment variables from '
                         'a .env file. This keeps passwords and secrets out of the code.'],
    ]
)

add_page_break()

# ============================================================================
# 3. FOLDER STRUCTURE
# ============================================================================

add_heading('3. How the Project is Organized (Folder Structure)', level=1)

add_para('The project has three main parts: frontend, backend, and database. Here is the structure:')

add_code('''CareerBridge/
|
|-- backend/                    <- PHP REST API (server-side code)
|   |-- .env                    <- Secret settings (database password, JWT key)
|   |-- .env.example            <- Template for .env (safe to share)
|   |-- .htaccess               <- Apache config (redirects all URLs to index.php)
|   |-- composer.json           <- PHP dependencies list (like package.json for PHP)
|   |-- composer.lock           <- Exact versions of PHP dependencies
|   |-- vendor/                 <- Installed PHP libraries (auto-generated)
|   |-- public/
|   |   |-- index.php           <- MAIN ENTRY POINT - all routes are defined here
|   |-- src/
|       |-- config/
|       |   |-- database.php    <- Database connection (PDO singleton)
|       |-- Middleware/
|       |   |-- CorsMiddleware.php   <- Handles CORS (cross-origin requests)
|       |   |-- JwtMiddleware.php    <- Verifies JWT tokens and checks roles
|       |-- Support/
|       |   |-- Json.php        <- Helper to write JSON responses
|       |-- Modules/            <- One folder per module
|           |-- Auth/AuthController.php           (M1)
|           |-- Jobs/JobController.php            (M2)
|           |-- Applications/ApplicationController.php (M3)
|           |-- Companies/CompanyController.php   (M4)
|           |-- Profile/ProfileController.php     (M5)
|           |-- Admin/AdminController.php         (M5+M6)
|           |-- Forums/ForumController.php        (M7)
|           |-- Interviews/InterviewController.php (M8)
|
|-- frontend/                   <- Vue.js SPA (client-side code)
|   |-- index.html              <- HTML shell that loads the Vue app
|   |-- package.json            <- JavaScript dependencies list
|   |-- vite.config.js          <- Vite config (dev server + API proxy)
|   |-- tailwind.config.js      <- Tailwind CSS configuration
|   |-- postcss.config.js       <- PostCSS config (processes Tailwind)
|   |-- src/
|   |   |-- main.js             <- App startup (creates Vue app, adds Pinia + Router)
|   |   |-- App.vue             <- Root component (NavBar + page content + Toast)
|   |   |-- style.css           <- Global styles + Tailwind imports + custom classes
|   |   |-- assets/             <- Static files (logo.svg)
|   |   |-- components/         <- Reusable UI components
|   |   |   |-- NavBar.vue          <- Top navigation bar (changes by role)
|   |   |   |-- JobCard.vue          <- Job listing card
|   |   |   |-- StatusBadge.vue      <- Colored status pill
|   |   |   |-- LoadingSpinner.vue   <- Spinning loader
|   |   |   |-- ToastNotification.vue <- Pop-up messages
|   |   |   |-- ConfirmDialog.vue     <- "Are you sure?" modal
|   |   |-- composables/
|   |   |   |-- useToast.js     <- Toast notification helper
|   |   |-- router/
|   |   |   |-- index.js        <- Route definitions + navigation guards
|   |   |-- services/
|   |   |   |-- api.js          <- Axios instance with JWT interceptor
|   |   |-- stores/
|   |   |   |-- auth.js         <- Pinia store for authentication state
|   |   |-- views/              <- Page components (one per page)
|   |       |-- LoginView.vue
|   |       |-- RegisterView.vue
|   |       |-- StudentDashboard.vue
|   |       |-- BrowseJobs.vue
|   |       |-- JobDetail.vue
|   |       |-- MyApplications.vue
|   |       |-- StudentProfile.vue
|   |       |-- AdminDashboard.vue
|   |       |-- ManageJobs.vue
|   |       |-- ManageApplications.vue
|   |       |-- ManageCompanies.vue
|   |       |-- AdminUsers.vue
|   |       |-- ForumList.vue
|   |       |-- ForumPost.vue
|   |       |-- Interviews.vue
|   |       |-- ManageInterviews.vue
|   |-- tests/                  <- Playwright E2E tests
|
|-- database/
|   |-- schema.sql              <- Creates all tables (run first)
|   |-- seed.sql                <- Inserts demo data (run second)
|
|-- docs/                       <- Documentation files
|-- git-schedule/               <- Per-member weekly commit plans
|-- README.md                   <- Setup instructions
''')

add_page_break()

# ============================================================================
# 4. HOW FRONTEND, BACKEND, DATABASE CONNECT
# ============================================================================

add_heading('4. How the Frontend, Backend, and Database Connect', level=1)

add_para('Here is the flow of how everything works together, explained step by step:')

add_number('The user opens the website in their browser at http://localhost:5173')
add_number('The browser loads the Vue.js app (from Vite dev server)')
add_number('The Vue app shows the login page')
add_number('The user types their email and password and clicks "Sign in"')
add_number('Vue sends a POST request to /api/auth/login using Axios')
add_number('Vite proxy forwards the request to the PHP backend at http://localhost:8000')
add_number('The PHP backend (Slim 4) receives the request in public/index.php')
add_number('The router matches the URL /api/auth/login to AuthController::login')
add_number('AuthController connects to MySQL using PDO (Database::getConnection())')
add_number('AuthController checks the users table for the email and verifies the password hash')
add_number('If correct, AuthController creates a JWT token and returns it as JSON')
add_number('The Vue frontend receives the token and stores it in localStorage')
add_number('For all future requests, Axios automatically attaches the JWT token')
add_number('When the user visits a protected page, JwtMiddleware checks the token')
add_number('If the token is valid and the role matches, the request proceeds')
add_number('The backend reads/writes data from MySQL and returns JSON responses')
add_number('Vue receives the JSON data and displays it on the page')

add_para('')
add_para('Simple diagram:', bold=True)
add_code('''
  Browser (User)
      |
      v
  Vue.js Frontend (localhost:5173)
      |
      | Axios HTTP requests (with JWT token)
      v
  Vite Proxy (forwards /api/* to backend)
      |
      v
  PHP Backend - Slim 4 (localhost:8000)
      |
      | PDO prepared statements
      v
  MySQL Database (careerbridge)
''')

add_page_break()

# ============================================================================
# 5. DATABASE SETUP
# ============================================================================

add_heading('5. Database Setup - How to Create and Fill the Database', level=1)

add_para('The database is MySQL. We have two SQL files that set everything up:')

add_para('Step 1: Create the database and tables', bold=True)
add_para('Run schema.sql first. This creates the database called "careerbridge" and all the tables:')
add_code('mysql -u root < database/schema.sql')

add_para('')
add_para('Step 2: Insert demo data (seed data)', bold=True)
add_para('Run seed.sql second. This inserts sample users, companies, jobs, applications, forum posts, etc.:')
add_code('mysql -u root < database/seed.sql')

add_para('')
add_para('What the seed data includes:', bold=True)
add_bullet('4 users: 1 super admin, 1 admin, 2 students (all with password: Password123!)')
add_bullet('3 companies: TechCorp, Axiata Digital, Petronas ICT')
add_bullet('10 job listings: mix of internships and full-time positions')
add_bullet('2 student profiles (for the 2 student users)')
add_bullet('5 job applications with different statuses (pending, reviewed, accepted, rejected)')
add_bullet('5 forum posts with tags, 6 comments, and post likes')
add_bullet('5 interview slots (4 open, 1 already booked)')
add_bullet('1 mock interview booking (pending)')

add_para('')
add_para('Important: The seed data uses dynamic dates (DATE_ADD(CURDATE(), INTERVAL 90 DAY))', bold=True)
add_para('This means deadlines and interview slots are always in the future when you run the seed. '
         'If you seeded a long time ago, re-run both SQL files to get fresh dates.')

add_para('')
add_para('How to re-seed (reset the database):', bold=True)
add_code('mysql -u root < database/schema.sql\nmysql -u root < database/seed.sql')

add_para('')
add_para('Note: schema.sql drops all tables first (DROP TABLE IF EXISTS), so it is safe to re-run. '
         'seed.sql truncates all tables before inserting, so it is also safe to re-run.')

add_page_break()

# ============================================================================
# 6. DATABASE TABLES
# ============================================================================

add_heading('6. Database Tables Explained (Schema)', level=1)

add_para('The database has 10 tables. Here is what each one stores:')

add_table(
    ['Table Name', 'Purpose', 'Key Columns', 'Related To'],
    [
        ['users', 'All user accounts (students, admins, super admins)',
         'user_id (PK), full_name, email, password_hash, role',
         'Referenced by almost every other table'],
        ['companies', 'Companies that post jobs',
         'company_id (PK), name, industry, location, description, created_by',
         'created_by -> users(user_id)'],
        ['jobs', 'Job and internship listings',
         'job_id (PK), company_id, posted_by, title, type, description, deadline, is_active',
         'company_id -> companies, posted_by -> users'],
        ['student_profiles', 'Extra profile info for students',
         'profile_id (PK), user_id, matric_no, programme, cgpa, skills, resume_text',
         'user_id -> users (1:1 relationship)'],
        ['applications', 'Job applications submitted by students',
         'application_id (PK), job_id, user_id, cover_letter, status, applied_at',
         'job_id -> jobs, user_id -> users. Unique: (job_id, user_id)'],
        ['posts', 'Forum discussion posts',
         'post_id (PK), user_id, title, content, tag, created_at',
         'user_id -> users'],
        ['comments', 'Comments on forum posts',
         'comment_id (PK), post_id, user_id, content, created_at',
         'post_id -> posts (CASCADE DELETE), user_id -> users'],
        ['post_likes', 'Likes on forum posts (one like per user per post)',
         'like_id (PK), post_id, user_id, created_at',
         'post_id -> posts (CASCADE DELETE), user_id -> users. Unique: (post_id, user_id)'],
        ['interview_slots', 'Time slots for mock interviews created by admins',
         'slot_id (PK), interviewer_id, scheduled_at, is_booked',
         'interviewer_id -> users'],
        ['mock_interviews', 'Student bookings for mock interviews + scores/feedback',
         'interview_id (PK), slot_id, student_id, job_category, status, score, feedback_text',
         'slot_id -> interview_slots (CASCADE DELETE), student_id -> users. Unique: (slot_id)'],
    ]
)

add_para('')
add_para('Key database design concepts used:', bold=True)

add_bullet('Primary Key (PK): A unique identifier for each row (e.g., user_id, job_id)')
add_bullet('Foreign Key (FK): A column that references another table (e.g., jobs.company_id points to companies.company_id)')
add_bullet('CASCADE DELETE: When a post is deleted, its comments and likes are automatically deleted too')
add_bullet('UNIQUE constraint: Prevents duplicates (e.g., a student cannot apply to the same job twice)')
add_bullet('ENUM: A column that can only have specific values (e.g., role can only be student, admin, or superadmin)')
add_bullet('AUTO_INCREMENT: The ID number automatically increases for each new row')
add_bullet('TIMESTAMP DEFAULT CURRENT_TIMESTAMP: Automatically records when a row is created')
add_bullet('ON UPDATE CURRENT_TIMESTAMP: Automatically updates the timestamp when a row is modified')

add_para('')
add_para('How tables relate to each other (Entity Relationships):', bold=True)
add_code('''
users (1) ----< (M) jobs          (one user posts many jobs)
users (1) ----< (M) applications   (one student submits many applications)
users (1) ----< (1) student_profiles (one student has one profile)
companies (1) ----< (M) jobs      (one company has many jobs)
jobs (1) ----< (M) applications    (one job receives many applications)
users (1) ----< (M) posts          (one user writes many posts)
posts (1) ----< (M) comments       (one post has many comments)
posts (1) ----< (M) post_likes     (one post has many likes)
users (1) ----< (M) interview_slots (one admin creates many slots)
interview_slots (1) ----< (1) mock_interviews (one slot = one booking)
''')

add_page_break()

# ============================================================================
# 7. BACKEND EXPLANATION
# ============================================================================

add_heading('7. Backend Explanation (PHP / Slim 4 REST API)', level=1)

add_para(
    'The backend is the server-side code. It receives HTTP requests from the frontend, '
    'processes them, talks to the database, and returns JSON responses. '
    'We use PHP with the Slim 4 framework.'
)

# ---- 7a. Entry Point ----
add_heading('7a. Entry Point (public/index.php) - How Routes Work', level=2)

add_para('This is the main file that starts the entire backend. Here is what it does:')

add_number('Loads the autoloader (vendor/autoload.php) - this makes all PHP libraries available')
add_number('Loads environment variables from .env (database credentials, JWT secret)')
add_number('Creates a Slim 4 application instance')
add_number('Adds CORS middleware (allows the frontend to make requests)')
add_number('Adds error handling middleware (catches errors and returns JSON)')
add_number('Defines all API routes (URLs) and maps them to controller methods')
add_number('Runs the application (starts listening for requests)')

add_para('')
add_para('How routes work:', bold=True)
add_para('Each route maps an HTTP method + URL path to a PHP function. For example:')

add_code('''// When someone sends GET to /api/jobs, run JobController::index
$app->get('/api/jobs', [JobController::class, 'index']);

// When someone sends POST to /api/jobs, run JobController::create
// BUT only if they have a JWT token with role admin or superadmin
$app->post('/api/jobs', [JobController::class, 'create'])
    ->add(new JwtMiddleware(['admin', 'superadmin']));''')

add_para('')
add_para('The ->add(new JwtMiddleware(...)) part is middleware. It runs BEFORE the controller. '
         'It checks if the user is logged in and has the right role. If not, it returns 401 or 403.')

add_para('')
add_para('All routes defined in index.php:', bold=True)

add_table(
    ['Method', 'Path', 'Controller::Method', 'Who Can Access'],
    [
        ['POST', '/api/auth/register', 'AuthController::register', 'Anyone (no login needed)'],
        ['POST', '/api/auth/login', 'AuthController::login', 'Anyone (no login needed)'],
        ['GET', '/api/jobs', 'JobController::index', 'Anyone (public)'],
        ['GET', '/api/jobs/{id}', 'JobController::show', 'Anyone (public)'],
        ['POST', '/api/jobs', 'JobController::create', 'Admin / Super Admin'],
        ['PUT', '/api/jobs/{id}', 'JobController::update', 'Admin / Super Admin'],
        ['DELETE', '/api/jobs/{id}', 'JobController::delete', 'Admin / Super Admin'],
        ['GET', '/api/applications', 'ApplicationController::index', 'Any logged-in user'],
        ['POST', '/api/applications', 'ApplicationController::create', 'Student only'],
        ['PUT', '/api/applications/{id}/status', 'ApplicationController::updateStatus', 'Admin / Super Admin'],
        ['DELETE', '/api/applications/{id}', 'ApplicationController::delete', 'Student (own + pending only)'],
        ['GET', '/api/profile', 'ProfileController::show', 'Student only'],
        ['PUT', '/api/profile', 'ProfileController::upsert', 'Student only'],
        ['GET', '/api/companies', 'CompanyController::index', 'Any logged-in user'],
        ['POST', '/api/companies', 'CompanyController::create', 'Admin / Super Admin'],
        ['PUT', '/api/companies/{id}', 'CompanyController::update', 'Admin / Super Admin'],
        ['DELETE', '/api/companies/{id}', 'CompanyController::delete', 'Admin / Super Admin'],
        ['GET', '/api/admin/users', 'AdminController::users', 'Super Admin only'],
        ['GET', '/api/admin/analytics', 'AdminController::analytics', 'Admin / Super Admin'],
        ['GET', '/api/forums', 'ForumController::index', 'Any logged-in user'],
        ['GET', '/api/forums/{id}', 'ForumController::show', 'Any logged-in user'],
        ['POST', '/api/forums', 'ForumController::create', 'Student only'],
        ['PUT', '/api/forums/{id}', 'ForumController::update', 'Student (own post only)'],
        ['DELETE', '/api/forums/{id}', 'ForumController::delete', 'Student (own) or Admin'],
        ['POST', '/api/forums/{id}/like', 'ForumController::toggleLike', 'Any logged-in user'],
        ['POST', '/api/forums/{id}/comments', 'ForumController::createComment', 'Student only'],
        ['DELETE', '/api/forums/comments/{id}', 'ForumController::deleteComment', 'Student (own) or Admin'],
        ['GET', '/api/interviews/slots', 'InterviewController::listSlots', 'Any logged-in user'],
        ['POST', '/api/interviews/slots', 'InterviewController::createSlot', 'Admin / Super Admin'],
        ['DELETE', '/api/interviews/slots/{id}', 'InterviewController::deleteSlot', 'Admin / Super Admin'],
        ['GET', '/api/interviews/my-sessions', 'InterviewController::mySessions', 'Student only'],
        ['POST', '/api/interviews/bookings', 'InterviewController::bookSlot', 'Student only'],
        ['PUT', '/api/interviews/bookings/{id}', 'InterviewController::updateBooking', 'Student only'],
        ['GET', '/api/interviews/admin/manage', 'InterviewController::adminList', 'Admin / Super Admin'],
        ['PUT', '/api/interviews/admin/evaluate/{id}', 'InterviewController::evaluate', 'Admin / Super Admin'],
    ]
)

add_page_break()

# ---- 7b. Database Connection ----
add_heading('7b. Database Connection (config/database.php)', level=2)

add_para('This file creates a connection to MySQL. It uses the Singleton pattern, which means '
         'only one database connection is created and shared across the entire app.')

add_para('Key things to know:', bold=True)
add_bullet('It reads database settings from .env file (DB_HOST, DB_NAME, DB_USER, DB_PASS)')
add_bullet('Uses PDO (PHP Data Objects) - a standard way to talk to databases in PHP')
add_bullet('Sets ERRMODE_EXCEPTION - database errors throw exceptions (errors we can catch)')
add_bullet('Sets FETCH_ASSOC - results come back as associative arrays (column names as keys)')
add_bullet('Sets EMULATE_PREPARES to false - uses real prepared statements (more secure)')

add_para('')
add_para('What is PDO and why is it important?', bold=True)
add_para(
    'PDO is PHP\'s way of talking to databases safely. The most important feature is '
    '"prepared statements." Instead of putting user input directly into SQL queries '
    '(which is dangerous - this is how SQL injection attacks happen), PDO uses placeholders. '
    'The SQL query and the data are sent separately to the database, so the database '
    'knows the data is just data, not SQL commands.'
)

add_para('')
add_para('Example of safe query (what we do):', bold=True)
add_code('''// SAFE: The user input goes through :email placeholder
$stmt = $pdo->prepare('SELECT * FROM users WHERE email = :email');
$stmt->execute([':email' => $userInput]);

// DANGEROUS (what we DON'T do - SQL injection risk):
// $sql = "SELECT * FROM users WHERE email = '" . $userInput . "'";''')

add_page_break()

# ---- 7c. JWT Middleware ----
add_heading('7c. JWT Authentication (JwtMiddleware.php)', level=2)

add_para('This is the security guard of our API. It runs before any protected route.')

add_para('What it does step by step:', bold=True)
add_number('Looks for the "Authorization: Bearer <token>" header in the request')
add_number('If no header found, returns 401 Unauthorized')
add_number('Extracts the token (the string after "Bearer ")')
add_number('Reads the JWT secret from .env')
add_number('Decodes and verifies the token using firebase/php-jwt library')
add_number('Checks if the token has expired (exp claim)')
add_number('If the route requires specific roles, checks if the user\'s role matches')
add_number('If role does not match, returns 403 Forbidden')
add_number('If everything is OK, attaches the user info to the request and lets it continue')

add_para('')
add_para('How roles work in the middleware:', bold=True)
add_code('''// Any logged-in user can access:
->add(new JwtMiddleware())

// Only students can access:
->add(new JwtMiddleware('student'))

// Only admins and super admins can access:
->add(new JwtMiddleware(['admin', 'superadmin']))

// Only super admins can access:
->add(new JwtMiddleware('superadmin'))''')

add_para('')
add_para('What is JWT?', bold=True)
add_para(
    'JWT (JSON Web Token) is like a digital ID card. When you log in, the backend creates '
    'a token that contains your user ID, email, role, and an expiry time. This token is '
    'signed with a secret key so nobody can fake it. The frontend stores this token and '
    'sends it with every request. The backend verifies the token to know who you are.'
)

add_para('')
add_para('JWT Structure (3 parts separated by dots):', bold=True)
add_code('''header.payload.signature

Example:
eyJ0eXAiOiJKV1QiLCJhbGciOiJIUzI1NiJ9    <- Header (algorithm info)
.eyJ1c2VyX2lkIjo0LCJlbWFpbCI6...        <- Payload (user_id, email, role, exp)
.j_ljGTMfvhTutgzC_DGorW1GWB90sUr4JIAGAEt-zBY  <- Signature

The payload contains:
{
  "user_id": 4,
  "email": "samin@student.utm.my",
  "role": "student",
  "full_name": "Samin Sarwat",
  "iat": 1782398221,    <- issued at (timestamp)
  "exp": 1782484621     <- expires at (24 hours later)
}''')

add_page_break()

# ---- 7d. CORS ----
add_heading('7d. CORS (CorsMiddleware.php)', level=2)

add_para(
    'CORS stands for Cross-Origin Resource Sharing. Browsers have a security rule: '
    'a website at one domain cannot make requests to a different domain by default. '
    'Our frontend runs at localhost:5173 and our backend at localhost:8000 - these are '
    'different "origins," so we need CORS.'
)

add_para('What our CORS middleware does:', bold=True)
add_bullet('Allows requests from localhost:5173 (our Vue dev server)')
add_bullet('Allows all HTTP methods (GET, POST, PUT, DELETE, OPTIONS)')
add_bullet('Allows Content-Type and Authorization headers')
add_bullet('Handles OPTIONS "preflight" requests (browsers send these before real requests)')
add_bullet('Can be opened to all origins with CORS_ALLOW_ALL=1 in .env (for testing)')

add_page_break()

# ---- 7e. JSON Helper ----
add_heading('7e. JSON Response Helper (Support/Json.php)', level=2)

add_para(
    'This is a small helper class that makes sure all our API responses look the same. '
    'Every response has a consistent format (called an "envelope"):'
)

add_code('''{
  "success": true,       <- true if the request worked, false if it failed
  "message": "...",      <- a human-readable message
  "data": { ... },       <- the actual data (only on success)
  "errors": { ... }      <- field-level validation errors (only on validation failure)
}''')

add_para('This makes it easy for the frontend to check if a request succeeded and display appropriate messages.')

add_page_break()

# ---- 7f. Auth Controller ----
add_heading('7f. Module 1: Auth Controller (Login & Register) - Owner: Areeb', level=2)

add_para('This controller handles user registration and login. It has two methods:')

add_para('register() - Create a new account:', bold=True)
add_number('Reads JSON body (full_name, email, password, role)')
add_number('Validates: name not empty, valid email, password >= 8 chars, role is student or admin')
add_number('Checks if email already exists in the database')
add_number('Hashes the password using PASSWORD_BCRYPT (never store plain text!)')
add_number('Inserts the new user into the users table')
add_number('Creates a JWT token and returns it with the user info')

add_para('')
add_para('login() - Sign in an existing user:', bold=True)
add_number('Reads JSON body (email, password)')
add_number('Validates: valid email, password not empty')
add_number('Looks up the user by email in the database')
add_number('Verifies the password using password_verify() (compares input to hash)')
add_number('If email or password is wrong, returns 401 Unauthorized')
add_number('If correct, creates a JWT token and returns it with the user info')

add_para('')
add_para('Key security points:', bold=True)
add_bullet('Passwords are hashed with bcrypt (cost 12) - even if someone steals the database, they cannot see passwords')
add_bullet('The JWT token expires after 24 hours')
add_bullet('The JWT is signed with a secret key in .env - nobody can forge a token without the secret')
add_bullet('Email is converted to lowercase to prevent duplicate accounts with different cases')

add_page_break()

# ---- 7g. Job Controller ----
add_heading('7g. Module 2: Job Controller (CRUD for Jobs) - Owner: Samin', level=2)

add_para('This controller manages job listings. It has full CRUD (Create, Read, Update, Delete).')

add_table(
    ['Method', 'HTTP', 'What It Does'],
    [
        ['index()', 'GET /api/jobs', 'Lists all active jobs. Supports search (by title or company name) '
                     'and type filter (internship/fulltime/parttime). Uses JOIN to get company info.'],
        ['show()', 'GET /api/jobs/{id}', 'Gets a single job by ID, with company name and poster name. '
                    'Returns 404 if not found.'],
        ['create()', 'POST /api/jobs', 'Creates a new job. Requires admin/superadmin. Validates all fields. '
                     'Sets is_active to 1 (visible). Returns the created job with company name.'],
        ['update()', 'PUT /api/jobs/{id}', 'Updates an existing job. Requires admin/superadmin. '
                     'Checks if job exists first (404 if not). Returns the updated job.'],
        ['delete()', 'DELETE /api/jobs/{id}', 'Soft-deletes a job (sets is_active = 0, does NOT remove from database). '
                     'This keeps the job record for audit purposes but hides it from students.'],
    ]
)

add_para('')
add_para('Validation in JobController:', bold=True)
add_bullet('company_id: required, must be a number')
add_bullet('title: required, max 150 characters')
add_bullet('type: must be internship, fulltime, or parttime')
add_bullet('description: required')
add_bullet('deadline: required, must be YYYY-MM-DD format, must be a future date')

add_para('')
add_para('Important: DELETE is a soft delete. The job stays in the database but is_active is set to 0. '
         'The index() method only returns jobs WHERE is_active = 1, so deleted jobs are hidden.')

add_page_break()

# ---- 7h. Application Controller ----
add_heading('7h. Module 3: Application Controller (Apply & Track) - Owner: Samin', level=2)

add_para('This controller handles job applications. Students apply, admins manage status.')

add_table(
    ['Method', 'HTTP', 'What It Does'],
    [
        ['index()', 'GET /api/applications', 'Lists applications. Students see only their own. '
                     'Admins see all. Supports status filter. Uses JOINs to show job title, '
                     'company name, and applicant info.'],
        ['create()', 'POST /api/applications', 'Student applies for a job. Checks: job exists, job is active, '
                     'student has not already applied (unique constraint). Returns 409 if duplicate.'],
        ['updateStatus()', 'PUT /api/applications/{id}/status', 'Admin changes application status. '
                     'Status must be: pending, reviewed, accepted, or rejected.'],
        ['delete()', 'DELETE /api/applications/{id}', 'Student withdraws their application. '
                     'Can only withdraw own applications. Can only withdraw if status is "pending". '
                     'This is a hard delete (row is removed from database).'],
    ]
)

add_para('')
add_para('Key business rules:', bold=True)
add_bullet('A student cannot apply to the same job twice (enforced by unique constraint in DB and checked in code)')
add_bullet('A student can only withdraw pending applications (once reviewed/accepted/rejected, cannot withdraw)')
add_bullet('A student can only withdraw their own applications (checks user_id matches)')
add_bullet('Admins see all applications; students see only their own (role-aware query)')

add_page_break()

# ---- 7i. Company Controller ----
add_heading('7i. Module 4: Company Controller (CRUD for Companies) - Owner: Mariam', level=2)

add_para('This controller manages the company directory. Full CRUD.')

add_table(
    ['Method', 'HTTP', 'What It Does'],
    [
        ['index()', 'GET /api/companies', 'Lists all companies, ordered by name. Any logged-in user can see.'],
        ['create()', 'POST /api/companies', 'Creates a new company. Admin/superadmin only. '
                     'Records who created it (created_by).'],
        ['update()', 'PUT /api/companies/{id}', 'Updates company info. Admin/superadmin only.'],
        ['delete()', 'DELETE /api/companies/{id}', 'Deletes a company. BUT: if the company still has jobs '
                     'attached, returns 400 error. This prevents orphaned jobs.'],
    ]
)

add_para('')
add_para('Validation: name (required, max 150), industry (required, max 100), location (required, max 150)')

add_page_break()

# ---- 7j. Profile Controller ----
add_heading('7j. Module 5: Profile Controller (Student Profiles) - Owner: Mariam', level=2)

add_para('This controller lets students create and update their profile (matric number, programme, CGPA, skills, resume).')

add_table(
    ['Method', 'HTTP', 'What It Does'],
    [
        ['show()', 'GET /api/profile', 'Gets the logged-in student\'s profile. Returns 404 if not yet created.'],
        ['upsert()', 'PUT /api/profile', 'Creates or updates the profile. Uses "INSERT ... ON DUPLICATE KEY UPDATE" '
                     'which means: if a profile already exists for this user, update it; if not, create it. '
                     'This is called an "upsert" (update + insert).'],
    ]
)

add_para('')
add_para('Validation: matric_no (required, unique, max 20), programme (required, max 100), '
         'cgpa (optional, must be 0.00-4.00)')

add_page_break()

# ---- 7k. Admin Controller ----
add_heading('7k. Module 5+6: Admin Controller (Analytics & Users) - Owners: Mariam & Areeb', level=2)

add_para('This controller has two methods for admin features:')

add_para('analytics() - GET /api/admin/analytics:', bold=True)
add_para('Returns dashboard statistics for the admin dashboard:')
add_bullet('total_jobs: count of all jobs')
add_bullet('active_jobs: count of jobs where is_active = 1')
add_bullet('total_applications: count of all applications')
add_bullet('pending_count, reviewed_count, accepted_count, rejected_count: counts by status')
add_bullet('applications_this_week: count of applications submitted this week')

add_para('')
add_para('users() - GET /api/admin/users:', bold=True)
add_para('Returns all user accounts (without password hashes). Super admin only. '
         'Used on the Admin Users page to see who is registered in the system.')

add_page_break()

# ---- 7l. Forum Controller ----
add_heading('7l. Module 7: Forum Controller (Posts, Comments, Likes) - Owner: Monika', level=2)

add_para('This is the most complex controller. It handles forum posts, comments, and likes.')

add_para('Posts:', bold=True)
add_bullet('index() - List all posts with search and tag filter. Sorted by likes (most liked first), then by date.')
add_bullet('show() - Get a single post with all its comments')
add_bullet('create() - Create a post (students only). Tag is free text, defaults to "General"')
add_bullet('update() - Edit your own post (checks user_id matches)')
add_bullet('delete() - Delete a post. Owner or admin can delete. CASCADE delete removes comments and likes automatically')

add_para('')
add_para('Likes:', bold=True)
add_bullet('toggleLike() - Like or unlike a post. If already liked, removes the like. If not liked, adds it.')
add_bullet('This is called a "toggle" - like a switch that turns on and off')
add_bullet('Each user can only like a post once (enforced by unique constraint in post_likes table)')

add_para('')
add_para('Comments:', bold=True)
add_bullet('createComment() - Add a comment to a post (students only)')
add_bullet('deleteComment() - Delete a comment. Owner or admin can delete.')
add_bullet('Comments are "flat" - no replies to replies (no nested comments)')

add_para('')
add_para('Special features:', bold=True)
add_bullet('The index() method also returns which posts the current user has liked (liked_by_me field)')
add_bullet('Post count and like count are calculated using subqueries in SQL')
add_bullet('Route ordering matters: /api/forums/comments/{id} is defined BEFORE /api/forums/{id} '
           'so Slim does not capture "comments" as a post ID')

add_page_break()

# ---- 7m. Interview Controller ----
add_heading('7m. Module 8: Interview Controller (Slots & Bookings) - Owner: Monika', level=2)

add_para('This controller manages mock interview scheduling. Admins create time slots, students book them.')

add_para('Slot Management (Admin):', bold=True)
add_bullet('listSlots() - List open slots (future + unbooked). Admins can see all with ?all=1')
add_bullet('createSlot() - Create a new interview slot. Validates: future date, interviewer must be admin/superadmin')
add_bullet('deleteSlot() - Delete a slot. Cannot delete if already booked (returns 409)')

add_para('')
add_para('Student Bookings:', bold=True)
add_bullet('mySessions() - View own bookings with score and feedback')
add_bullet('bookSlot() - Book an available slot. Uses database transaction + row locking (FOR UPDATE) '
           'to prevent two students booking the same slot at the same time')
add_bullet('updateBooking() - Change job category or cancel. Cancelling frees the slot for others')

add_para('')
add_para('Admin Evaluation:', bold=True)
add_bullet('adminList() - View all bookings in the system')
add_bullet('evaluate() - Submit a score (0-100) and feedback text for a completed interview')

add_para('')
add_para('Key technical feature - Row Locking:', bold=True)
add_para(
    'When a student books a slot, the code uses "SELECT ... FOR UPDATE" inside a transaction. '
    'This locks the slot row so if two students try to book the same slot at the exact same time, '
    'one has to wait for the other. The first one gets the booking, the second one gets a 409 error. '
    'This prevents "double booking" - a common race condition problem.'
)

add_code('''// Transaction with row locking (prevents double-booking)
$pdo->beginTransaction();
$stmt = $pdo->prepare(
    'SELECT slot_id, is_booked, scheduled_at
     FROM interview_slots
     WHERE slot_id = :id
     LIMIT 1
     FOR UPDATE'    // <-- This locks the row
);
// ... check if already booked, insert booking, update slot ...
$pdo->commit();''')

add_page_break()

# ============================================================================
# 8. FRONTEND EXPLANATION
# ============================================================================

add_heading('8. Frontend Explanation (Vue.js 3 SPA)', level=1)

add_para(
    'The frontend is a Single Page Application (SPA) built with Vue.js 3. '
    'In a SPA, the browser loads one HTML page, and JavaScript handles all navigation. '
    'Pages change without reloading the browser, making the app feel fast and smooth.'
)

# ---- 8a. Main Entry ----
add_heading('8a. Main Entry (main.js & App.vue)', level=2)

add_para('main.js - The startup file:', bold=True)
add_number('Creates a Vue app instance')
add_number('Adds Pinia (for state management)')
add_number('Restores the login session from localStorage (so you stay logged in on refresh)')
add_number('Adds the Vue Router (for page navigation)')
add_number('Mounts the app to the #app div in index.html')

add_para('')
add_para('App.vue - The root component:', bold=True)
add_bullet('Shows the NavBar at the top (only when logged in)')
add_bullet('Shows the current page content using <RouterView>')
add_bullet('Shows toast notifications (pop-up messages)')
add_bullet('Has a fade transition between pages for smooth navigation')

add_page_break()

# ---- 8b. Router ----
add_heading('8b. Router (How Pages Navigate)', level=2)

add_para('The router defines all the pages (routes) in the app and controls who can see them.')

add_para('Route definitions:', bold=True)
add_table(
    ['Path', 'Page', 'Who Can See'],
    [
        ['/login', 'LoginView', 'Public (not logged in only)'],
        ['/register', 'RegisterView', 'Public (not logged in only)'],
        ['/student/dashboard', 'StudentDashboard', 'Students'],
        ['/student/jobs', 'BrowseJobs', 'Students'],
        ['/student/jobs/:id', 'JobDetail', 'Students'],
        ['/student/applications', 'MyApplications', 'Students'],
        ['/student/profile', 'StudentProfile', 'Students'],
        ['/admin/dashboard', 'AdminDashboard', 'Admins + Super Admins'],
        ['/admin/jobs', 'ManageJobs', 'Admins + Super Admins'],
        ['/admin/applications', 'ManageApplications', 'Admins + Super Admins'],
        ['/admin/companies', 'ManageCompanies', 'Admins + Super Admins'],
        ['/admin/interviews', 'ManageInterviews', 'Admins + Super Admins'],
        ['/admin/users', 'AdminUsers', 'Super Admins only'],
        ['/forum', 'ForumList', 'All logged-in users'],
        ['/forum/:id', 'ForumPost', 'All logged-in users'],
        ['/interviews', 'Interviews', 'Students'],
    ]
)

add_para('')
add_para('Navigation Guards (route protection):', bold=True)
add_para('Before each navigation, the router checks:')
add_number('If the page is public (login/register) and user is already logged in -> redirect to dashboard')
add_number('If the page is protected and user is not logged in -> redirect to login page')
add_number('If the page requires a specific role and user does not have it -> redirect to their dashboard')
add_number('If all checks pass -> allow navigation')

add_para('')
add_para('roleHome() function:', bold=True)
add_para('When a user visits "/" (root), they are redirected based on their role:')
add_bullet('Not logged in -> /login')
add_bullet('Admin/Super Admin -> /admin/dashboard')
add_bullet('Student -> /student/dashboard')

add_page_break()

# ---- 8c. Auth Store ----
add_heading('8c. Auth Store (Pinia - Login State Management)', level=2)

add_para('The auth store manages who is logged in. It uses Pinia (Vue\'s state management library).')

add_para('What it stores:', bold=True)
add_bullet('token: the JWT token (null if not logged in)')
add_bullet('user: the user object { user_id, email, role, full_name } (null if not logged in)')

add_para('')
add_para('Getters (computed properties):', bold=True)
add_bullet('isAuthenticated: true if token exists')
add_bullet('isStudent: true if role is "student"')
add_bullet('isAdmin: true if role is "admin" or "superadmin"')
add_bullet('isSuperAdmin: true if role is "superadmin"')

add_para('')
add_para('Actions (methods):', bold=True)
add_bullet('initFromStorage(): Called on app startup. Reads token from localStorage. '
           'Decodes the JWT to get user info. Checks if token is expired.')
add_bullet('login(email, password): Sends POST to /api/auth/login. Stores token and user. '
           'Saves to localStorage so login persists across page refreshes.')
add_bullet('register(payload): Sends POST to /api/auth/register. Auto-logs in the new user.')
add_bullet('logout(): Clears token and user from state and localStorage. Redirects to /login.')

add_para('')
add_para('How the JWT is decoded on the frontend:', bold=True)
add_para(
    'The store has a decodeToken() function that reads the middle part of the JWT '
    '(the payload), base64-decodes it, and extracts the user info. This is only for '
    'display purposes - the backend is responsible for actual verification.'
)

add_page_break()

# ---- 8d. API Service ----
add_heading('8d. API Service (Axios - How Frontend Talks to Backend)', level=2)

add_para('This file creates a shared Axios instance used by every view and store.')

add_para('What it does:', bold=True)
add_bullet('Sets baseURL to "/api" (all requests go to /api/... which Vite proxies to the backend)')
add_bullet('Sets timeout to 15 seconds')
add_bullet('Request interceptor: automatically attaches the JWT token from localStorage to every request')
add_bullet('Response interceptor: if any request returns 401 (unauthorized), clears the token and redirects to login')

add_code('''// Request interceptor - adds JWT token to every request
api.interceptors.request.use((config) => {
  const token = localStorage.getItem('cb_token')
  if (token) {
    config.headers.Authorization = `Bearer ${token}`
  }
  return config
})

// Response interceptor - auto logout on 401
api.interceptors.response.use(
  (response) => response,
  async (error) => {
    if (error.response && error.response.status === 401) {
      localStorage.removeItem('cb_token')
      localStorage.removeItem('cb_user')
      router.push('/login')
    }
    return Promise.reject(error)
  }
)''')

add_para('')
add_para('Vite Proxy (vite.config.js):', bold=True)
add_para(
    'During development, the Vue dev server runs on port 5173 and the PHP backend on port 8000. '
    'The Vite proxy forwards any request starting with /api to http://localhost:8000. '
    'This avoids CORS issues during development and makes the frontend code simpler '
    '(it just calls /api/jobs instead of http://localhost:8000/api/jobs).'
)

add_page_break()

# ---- 8e. Components ----
add_heading('8e. Components (Reusable UI Parts)', level=2)

add_para('Components are reusable pieces of UI that are used across multiple pages:')

add_table(
    ['Component', 'What It Does'],
    [
        ['NavBar.vue', 'Top navigation bar. Shows different links based on user role. '
                       'Has desktop and mobile (hamburger menu) versions. Shows user name and logout button.'],
        ['JobCard.vue', 'A card showing a job listing (title, company, type badge, short description, deadline). '
                        'Has "View" and "Apply" buttons. Used on the Browse Jobs page.'],
        ['StatusBadge.vue', 'A colored pill showing application status. Colors: pending=amber, '
                            'reviewed=blue, accepted=green, rejected=red.'],
        ['LoadingSpinner.vue', 'A spinning circle with "Loading..." text. Shown while waiting for API responses.'],
        ['ToastNotification.vue', 'Pop-up messages that appear in the top-right corner. '
                                  'Green for success, red for error, blue for info. Auto-dismisses after 3 seconds.'],
        ['ConfirmDialog.vue', 'A "Are you sure?" modal dialog. Used before destructive actions '
                              '(delete, withdraw, deactivate). Supports keyboard (Escape to cancel, Enter to confirm).'],
    ]
)

add_para('')
add_para('useToast.js (Composable):', bold=True)
add_para(
    'A composable is a reusable function in Vue 3. useToast provides a way for any view '
    'to show toast notifications. It uses a reactive state object that the ToastNotification '
    'component reads from. Any view can call useToast().success("message") to show a green toast.'
)

add_page_break()

# ---- 8f. Views ----
add_heading('8f. Views (Page-by-Page Explanation)', level=2)

add_para('Each view is a page in the application. Here is what each one does:')

add_table(
    ['View', 'Page', 'What It Does'],
    [
        ['LoginView.vue', 'Login page',
         'Email and password form. Client-side validation (valid email, password >= 8 chars). '
         'On submit, calls auth store login(). Redirects to dashboard on success.'],
        ['RegisterView.vue', 'Register page',
         'New account form (name, email, password, confirm password, role selection). '
         'Client-side validation. On submit, calls auth store register(). Auto-logs in.'],
        ['StudentDashboard.vue', 'Student dashboard',
         'Shows welcome message, stat cards (total applied, pending, accepted, rejected), '
         'and a table of 5 most recent applications with status badges.'],
        ['BrowseJobs.vue', 'Browse jobs (student)',
         'Lists all active jobs as cards. Has search box (with 400ms debounce) and type filter dropdown. '
         'Apply button opens a modal with a cover letter textarea. Calls POST /api/applications.'],
        ['JobDetail.vue', 'Job detail (student)',
         'Shows full job info (title, company, type, description, requirements, deadline). '
         'Has an apply button that opens the same apply modal.'],
        ['MyApplications.vue', 'My applications (student)',
         'Lists student\'s applications with tabs (All, Pending, Reviewed, Accepted, Rejected). '
         'Can withdraw pending applications (with confirm dialog).'],
        ['StudentProfile.vue', 'Student profile',
         'Form to create/edit student profile (matric no, programme, CGPA, skills, resume text). '
         'Uses PUT /api/profile (upsert). Shows 404 gracefully if profile not yet created.'],
        ['AdminDashboard.vue', 'Admin dashboard',
         'Shows stat cards (active jobs, total applications, pending, accepted, rejected). '
         'Has a Chart.js bar chart showing application status breakdown. '
         'Quick action buttons to post jobs or view applications.'],
        ['ManageJobs.vue', 'Manage jobs (admin)',
         'Table of all jobs with edit and delete (deactivate) buttons. '
         'Create/Edit modal with form (title, company dropdown, type, deadline, description, requirements). '
         'Client-side validation. Uses ConfirmDialog before deactivating.'],
        ['ManageApplications.vue', 'Manage applications (admin)',
         'Table of ALL applications with tabs by status. '
         'Admin can change status using a dropdown (pending/reviewed/accepted/rejected). '
         'Uses optimistic update (updates UI immediately, rolls back on error).'],
        ['ManageCompanies.vue', 'Manage companies (admin)',
         'Table of companies with edit and delete buttons. '
         'Create/Edit modal with form (name, industry, location, description). '
         'Delete is blocked if company has jobs attached (shows error toast).'],
        ['AdminUsers.vue', 'Users (super admin)',
         'Read-only table of all user accounts. Has search by name or email. '
         'Shows user ID, name, email, role (with colored badge), and join date.'],
        ['ForumList.vue', 'Forum list',
         'Lists all forum posts sorted by likes. Has search box and tag filter. '
         'Each post shows title, author, tag, like count, comment count, and like button.'],
        ['ForumPost.vue', 'Forum post detail',
         'Shows full post content, like button, and comments list. '
         'Students can add comments. Post owner or admin can delete posts/comments. '
         'Students can edit their own posts.'],
        ['Interviews.vue', 'Mock interviews (student)',
         'Shows available interview slots (future, unbooked). '
         'Students can book a slot with a job category. '
         'Shows their own sessions with status, score, and feedback.'],
        ['ManageInterviews.vue', 'Manage interviews (admin)',
         'Admins can create interview slots (date/time, interviewer). '
         'View all bookings. Evaluate completed interviews (score 0-100 + feedback text). '
         'Can delete unbooked slots.'],
    ]
)

add_page_break()

# ============================================================================
# 9. JWT AUTHENTICATION FLOW
# ============================================================================

add_heading('9. JWT Authentication Flow (Step by Step)', level=1)

add_para('Here is the complete flow of how authentication works in our app:')

add_number('User goes to the login page and enters email + password')
add_number('Vue calls auth.login() which sends POST /api/auth/login with email and password')
add_number('Backend AuthController::login() receives the request')
add_number('Backend looks up the user by email in the users table')
add_number('Backend uses password_verify() to check the password against the stored hash')
add_number('If correct, backend creates a JWT with: user_id, email, role, full_name, issued time, expiry (24h)')
add_number('Backend signs the JWT with the secret key from .env (HS256 algorithm)')
add_number('Backend returns { success: true, data: { token: "...", user: {...} } }')
add_number('Frontend stores the token in localStorage as "cb_token"')
add_number('Frontend stores user info in localStorage as "cb_user"')
add_number('Frontend Pinia store updates its state (token and user)')
add_number('Router redirects user to their dashboard based on role')
add_number('For every future API request, Axios interceptor adds "Authorization: Bearer <token>" header')
add_number('Backend JwtMiddleware receives the request, extracts the token, verifies signature and expiry')
add_number('If token is valid, middleware attaches user info to the request and lets it continue')
add_number('If token is invalid/expired, middleware returns 401 and frontend redirects to login')
add_number('On page refresh, initFromStorage() reads the token from localStorage and restores the session')
add_number('On logout, token is removed from localStorage and Pinia state is cleared')

add_page_break()

# ============================================================================
# 10. CRUD SUMMARY
# ============================================================================

add_heading('10. CRUD Operations Summary (Where is GET, POST, PUT, DELETE?)', level=1)

add_para('CRUD stands for Create, Read, Update, Delete. The rubric requires all four operations. '
         'Here is where each one is implemented:')

add_table(
    ['Operation', 'HTTP Method', 'Module', 'Endpoint', 'Description'],
    [
        ['Create', 'POST', 'M1 Auth', '/api/auth/register', 'Create a new user account'],
        ['Create', 'POST', 'M2 Jobs', '/api/jobs', 'Create a new job listing'],
        ['Create', 'POST', 'M3 Applications', '/api/applications', 'Apply for a job'],
        ['Create', 'POST', 'M4 Companies', '/api/companies', 'Create a new company'],
        ['Create (Upsert)', 'PUT', 'M5 Profile', '/api/profile', 'Create or update student profile'],
        ['Create', 'POST', 'M7 Forums', '/api/forums', 'Create a forum post'],
        ['Create', 'POST', 'M7 Forums', '/api/forums/{id}/comments', 'Add a comment'],
        ['Create', 'POST', 'M8 Interviews', '/api/interviews/slots', 'Create an interview slot'],
        ['Create', 'POST', 'M8 Interviews', '/api/interviews/bookings', 'Book an interview slot'],
        ['Read', 'GET', 'M2 Jobs', '/api/jobs', 'List all active jobs (with search/filter)'],
        ['Read', 'GET', 'M2 Jobs', '/api/jobs/{id}', 'Get a single job'],
        ['Read', 'GET', 'M3 Applications', '/api/applications', 'List applications (role-aware)'],
        ['Read', 'GET', 'M4 Companies', '/api/companies', 'List all companies'],
        ['Read', 'GET', 'M5 Profile', '/api/profile', 'Get student profile'],
        ['Read', 'GET', 'M5 Admin', '/api/admin/analytics', 'Get dashboard statistics'],
        ['Read', 'GET', 'M6 Admin', '/api/admin/users', 'List all users'],
        ['Read', 'GET', 'M7 Forums', '/api/forums', 'List all forum posts'],
        ['Read', 'GET', 'M7 Forums', '/api/forums/{id}', 'Get a single post with comments'],
        ['Read', 'GET', 'M8 Interviews', '/api/interviews/slots', 'List interview slots'],
        ['Read', 'GET', 'M8 Interviews', '/api/interviews/my-sessions', 'Get student\'s bookings'],
        ['Update', 'PUT', 'M2 Jobs', '/api/jobs/{id}', 'Update a job listing'],
        ['Update', 'PUT', 'M3 Applications', '/api/applications/{id}/status', 'Change application status'],
        ['Update', 'PUT', 'M4 Companies', '/api/companies/{id}', 'Update company info'],
        ['Update', 'PUT', 'M5 Profile', '/api/profile', 'Update student profile (upsert)'],
        ['Update', 'PUT', 'M7 Forums', '/api/forums/{id}', 'Edit own post'],
        ['Update', 'PUT', 'M8 Interviews', '/api/interviews/bookings/{id}', 'Change category or cancel'],
        ['Update', 'PUT', 'M8 Interviews', '/api/interviews/admin/evaluate/{id}', 'Submit score + feedback'],
        ['Delete', 'DELETE', 'M2 Jobs', '/api/jobs/{id}', 'Soft-delete (deactivate) a job'],
        ['Delete', 'DELETE', 'M3 Applications', '/api/applications/{id}', 'Withdraw a pending application'],
        ['Delete', 'DELETE', 'M4 Companies', '/api/companies/{id}', 'Delete a company (if no jobs)'],
        ['Delete', 'DELETE', 'M7 Forums', '/api/forums/{id}', 'Delete a post (cascade)'],
        ['Delete', 'DELETE', 'M7 Forums', '/api/forums/comments/{id}', 'Delete a comment'],
        ['Delete', 'DELETE', 'M8 Interviews', '/api/interviews/slots/{id}', 'Delete an unbooked slot'],
    ]
)

add_page_break()

# ============================================================================
# 11. INPUT VALIDATION
# ============================================================================

add_heading('11. Input Validation (Frontend and Backend)', level=1)

add_para('The rubric requires input validation on both frontend and backend. Here is how we do it:')

add_para('Frontend Validation (Vue.js):', bold=True)
add_bullet('Login: email format check, password not empty, password >= 8 chars')
add_bullet('Register: name required (max 100), valid email, password >= 8 chars, passwords match, valid role')
add_bullet('Job form: title required (max 150), type must be valid, description required, deadline required and future')
add_bullet('Company form: name required (max 150), industry required (max 100), location required (max 150)')
add_bullet('Profile form: matric_no required (max 20), programme required (max 100), CGPA 0.00-4.00')
add_bullet('Forum post: title required (max 150), content required (max 5000), tag max 60')
add_bullet('Comment: content required (max 2000)')
add_bullet('Interview slot: scheduled_at required and future, job_category required (max 100)')
add_bullet('Evaluation: score 0-100, feedback required (max 5000)')

add_para('')
add_para('Backend Validation (PHP):', bold=True)
add_bullet('Every controller validates input before processing')
add_bullet('Same rules as frontend, but enforced server-side (cannot be bypassed)')
add_bullet('Returns 400 with field-level errors in the "errors" object')
add_bullet('Email format checked with PHP filter_var(FILTER_VALIDATE_EMAIL)')
add_bullet('Date format checked with DateTime::createFromFormat')
add_bullet('Numeric values checked with is_numeric() and range checks')

add_para('')
add_para('Why validate on both sides?', bold=True)
add_para(
    'Frontend validation gives instant feedback to the user (no waiting for server). '
    'But it can be bypassed (someone could use Postman or curl to send bad data directly to the API). '
    'Backend validation is the real security - it always runs and cannot be bypassed.'
)

add_page_break()

# ============================================================================
# 12. ERROR HANDLING
# ============================================================================

add_heading('12. Error Handling & HTTP Status Codes', level=1)

add_para('We use proper HTTP status codes to indicate the result of each request:')

add_table(
    ['Status Code', 'Meaning', 'When We Use It'],
    [
        ['200 OK', 'Request succeeded', 'Successful GET, PUT, DELETE'],
        ['201 Created', 'Resource created', 'Successful POST (create)'],
        ['400 Bad Request', 'Invalid input', 'Validation failed, invalid ID, bad data'],
        ['401 Unauthorized', 'Not logged in or invalid token', 'Missing/invalid JWT, wrong password'],
        ['403 Forbidden', 'Logged in but not allowed', 'Wrong role (e.g., student trying to create a job)'],
        ['404 Not Found', 'Resource does not exist', 'Job/application/post/slot not found'],
        ['409 Conflict', 'Duplicate or conflict', 'Already applied, slot already booked, delete booked slot'],
        ['500 Server Error', 'Something went wrong on server', 'Database errors, unhandled exceptions'],
    ]
)

add_para('')
add_para('Global error handler:', bold=True)
add_para(
    'In index.php, we set a custom error handler that catches all unhandled exceptions and '
    'returns a JSON response instead of showing a raw PHP error. In development mode '
    '(APP_DEBUG=1 in .env), it includes the error message. In production, it returns '
    'a generic "Server error" message to avoid leaking sensitive info.'
)

add_para('')
add_para('404 fallback:', bold=True)
add_para('Any URL that does not match a defined route returns a 404 JSON response: '
         '{ "success": false, "message": "Route not found." }')

add_page_break()

# ============================================================================
# 13. SECURITY
# ============================================================================

add_heading('13. Security Features', level=1)

add_para('Here are all the security measures we implemented:')

add_table(
    ['Security Feature', 'Where', 'How It Works'],
    [
        ['Password Hashing', 'AuthController.php',
         'Uses PASSWORD_BCRYPT (cost 12). Passwords are never stored in plain text. '
         'Even if someone steals the database, they cannot reverse the hash.'],
        ['JWT Authentication', 'JwtMiddleware.php',
         'Every protected route requires a valid JWT token. Tokens are signed with HS256 '
         'and expire after 24 hours.'],
        ['Role-Based Access Control', 'JwtMiddleware.php + index.php',
         'Each route specifies which roles can access it. Students cannot access admin routes. '
         'Only super admins can view all users.'],
        ['SQL Injection Prevention', 'All controllers',
         'Every database query uses PDO prepared statements with named placeholders (:name). '
         'User input is NEVER concatenated into SQL strings.'],
        ['CORS Protection', 'CorsMiddleware.php',
         'Only allows requests from our frontend (localhost:5173). '
         'Can be restricted further by not setting CORS_ALLOW_ALL.'],
        ['JWT Secret in .env', 'backend/.env',
         'The JWT signing key is stored in .env, which is in .gitignore and never committed to GitHub.'],
        ['Auto Logout on 401', 'frontend/src/services/api.js',
         'If any API request returns 401, the frontend automatically clears the token '
         'and redirects to the login page.'],
        ['Ownership Checks', 'Application, Forum, Interview controllers',
         'Students can only delete/edit their own applications, posts, comments, and bookings. '
         'The backend checks user_id matches before allowing modifications.'],
        ['Soft Delete for Jobs', 'JobController.php',
         'Jobs are not actually deleted from the database. is_active is set to 0. '
         'This preserves data integrity and audit trail.'],
        ['Referential Integrity', 'Database schema',
         'Foreign key constraints prevent orphaned records. Companies with jobs cannot be deleted. '
         'Cascade delete on posts removes comments and likes automatically.'],
        ['Input Validation', 'All controllers + Vue views',
         'Both frontend and backend validate all input. See section 11 for details.'],
        ['No Password in User List', 'AdminController.php',
         'The admin users endpoint explicitly excludes password_hash from the SELECT query.'],
    ]
)

add_page_break()

# ============================================================================
# 14. HOW TO RUN
# ============================================================================

add_heading('14. How to Run the Project (Step by Step)', level=1)

add_para('Prerequisites (must be installed on your computer):', bold=True)
add_bullet('PHP 7.4 or newer (we have PHP 8.5)')
add_bullet('Composer (PHP package manager)')
add_bullet('Node.js 18+ and npm (we have Node 23)')
add_bullet('MySQL 5.7+ or MariaDB 10.3+ (we have MySQL 9.2)')

add_para('')
add_para('Step 1: Clone the repository', bold=True)
add_code('git clone https://github.com/saminsarwat007/webtechnologyproject.git\ncd webtechnologyproject')

add_para('')
add_para('Step 2: Set up the database', bold=True)
add_code('mysql -u root < database/schema.sql\nmysql -u root < database/seed.sql')

add_para('')
add_para('Step 3: Install backend dependencies', bold=True)
add_code('cd backend\ncomposer install\ncp .env.example .env\n# Edit .env if your MySQL has a password\ncd ..')

add_para('')
add_para('Step 4: Install frontend dependencies', bold=True)
add_code('cd frontend\nnpm install\ncd ..')

add_para('')
add_para('Step 5: Start the backend (Terminal 1)', bold=True)
add_code('cd backend\nphp -S localhost:8000 -t public')

add_para('')
add_para('Step 6: Start the frontend (Terminal 2)', bold=True)
add_code('cd frontend\nnpm run dev')

add_para('')
add_para('Step 7: Open the website', bold=True)
add_code('Open browser to: http://localhost:5173')

add_para('')
add_para('Important notes:', bold=True)
add_bullet('You need TWO terminals running at the same time (backend + frontend)')
add_bullet('The backend must be running for the frontend to work (API calls will fail otherwise)')
add_bullet('If you change .env, restart the PHP server')
add_bullet('If you need fresh demo data, re-run schema.sql and seed.sql')

add_page_break()

# ============================================================================
# 15. DEMO CREDENTIALS
# ============================================================================

add_heading('15. Demo Login Credentials', level=1)

add_para('All seeded accounts use the same password: Password123!', bold=True)

add_table(
    ['Role', 'Email', 'What You Can Demo'],
    [
        ['Super Admin', 'superadmin@careerbridge.my',
         'Everything an admin can do, PLUS view all users on the Users page'],
        ['Admin', 'farizal@careerbridge.my',
         'Dashboard with charts, manage jobs, manage applications, manage companies, '
         'manage interviews, participate in forum'],
        ['Student', 'samin@student.utm.my',
         'Dashboard, browse and apply for jobs, track applications, manage profile, '
         'participate in forum, book mock interviews'],
        ['Student', 'areeb@student.utm.my',
         'Same as above - second student account for testing'],
    ]
)

add_para('')
add_para('Recommended demo flow for presentation:', bold=True)
add_number('Log in as student (samin@student.utm.my / Password123!)')
add_number('Show the student dashboard (stats + recent applications)')
add_number('Browse jobs - show search and type filter')
add_number('Click "Apply" on a job - show the apply modal with cover letter')
add_number('Go to "My Applications" - show tabs and status badges')
add_number('Go to "Forum" - show posts, like a post, open a post, add a comment')
add_number('Go to "Mock Interviews" - show available slots, book one')
add_number('Go to "Profile" - show the profile form')
add_number('Log out and log in as admin (farizal@careerbridge.my / Password123!)')
add_number('Show admin dashboard with Chart.js bar chart')
add_number('Go to "Manage Jobs" - create a new job, edit one, deactivate one')
add_number('Go to "Applications" - change an application status')
add_number('Go to "Companies" - show the company table')
add_number('Go to "Interviews" - show bookings, evaluate one (add score + feedback)')
add_number('Log out and log in as super admin (superadmin@careerbridge.my / Password123!)')
add_number('Show the "Users" page (only visible to super admin)')

add_page_break()

# ============================================================================
# 16. QUESTIONS THE LECTURER MIGHT ASK
# ============================================================================

add_heading('16. Questions the Lecturer Might Ask (With Answers)', level=1)

# Q1
add_para('Q: What is a Single Page Application (SPA) and why did you use Vue.js?', bold=True)
add_para(
    'A: A SPA is a web application that loads a single HTML page and dynamically updates it '
    'using JavaScript, without reloading the browser. We used Vue.js 3 because it is a '
    'progressive framework that is easy to learn, has excellent documentation, and provides '
    'reactive data binding, component-based architecture, and a router for SPA navigation. '
    'Vue Router handles page transitions without browser reloads, making the app feel fast '
    'and smooth like a desktop application.'
)

# Q2
add_para('Q: How does your JWT authentication work?', bold=True)
add_para(
    'A: When a user logs in, the backend creates a JWT (JSON Web Token) containing their '
    'user_id, email, role, and an expiry time (24 hours). This token is signed with HS256 '
    'using a secret key stored in .env. The frontend stores the token in localStorage and '
    'sends it in the Authorization header with every request. The backend JwtMiddleware '
    'verifies the token signature and expiry before allowing access to protected routes. '
    'If the token is invalid or expired, it returns 401 and the frontend redirects to login.'
)

# Q3
add_para('Q: How do you prevent SQL injection?', bold=True)
add_para(
    'A: Every database query in our backend uses PDO prepared statements with named placeholders. '
    'For example, instead of "SELECT * FROM users WHERE email = \'$email\'", we use '
    '"SELECT * FROM users WHERE email = :email" and then execute with [\'email\' => $email]. '
    'The SQL query and the data are sent separately to MySQL, so the database treats user input '
    'as data, not as SQL commands. This makes SQL injection impossible.'
)

# Q4
add_para('Q: What is CORS and why do you need it?', bold=True)
add_para(
    'A: CORS (Cross-Origin Resource Sharing) is a browser security feature. When our frontend '
    '(localhost:5173) makes requests to our backend (localhost:8000), the browser considers '
    'these as different "origins" and blocks them by default. Our CorsMiddleware adds headers '
    'that tell the browser it is OK to allow these requests. It also handles OPTIONS preflight '
    'requests that browsers send before the actual request.'
)

# Q5
add_para('Q: How does role-based access control work?', bold=True)
add_para(
    'A: We have three roles: student, admin, and superadmin. Each API route in index.php '
    'specifies which roles can access it using JwtMiddleware. For example, '
    'new JwtMiddleware([\'admin\', \'superadmin\']) means only admins and super admins can '
    'access that route. On the frontend, the Vue Router has navigation guards that check '
    'the user\'s role before allowing access to a page. If a student tries to visit /admin/dashboard, '
    'they are redirected to their own dashboard.'
)

# Q6
add_para('Q: What is the difference between a soft delete and a hard delete?', bold=True)
add_para(
    'A: A soft delete sets is_active to 0 but keeps the record in the database. We use this '
    'for jobs - when an admin "deletes" a job, it is just hidden from students but the data '
    'remains for audit purposes. A hard delete actually removes the row from the database. '
    'We use this for applications when a student withdraws - the application row is deleted.'
)

# Q7
add_para('Q: How does the forum like toggle work?', bold=True)
add_para(
    'A: We have a post_likes table with a unique constraint on (post_id, user_id). '
    'When a user clicks "like", the backend checks if a like already exists. If it does, '
    'the like is removed (DELETE). If it does not, a like is added (INSERT). The like count '
    'is calculated dynamically using COUNT(*) in a subquery, not stored as a number. '
    'This ensures the count is always accurate.'
)

# Q8
add_para('Q: How do you prevent double-booking in mock interviews?', bold=True)
add_para(
    'A: We use database transactions with row locking. When a student books a slot, '
    'the backend starts a transaction and runs "SELECT ... FOR UPDATE" on the slot row. '
    'This locks the row so if two students try to book the same slot at the same time, '
    'one has to wait. The first student gets the booking, the second one gets a 409 error. '
    'The transaction ensures both the booking insert and the slot update happen together '
    'or not at all.'
)

# Q9
add_para('Q: What is Pinia and why do you use it?', bold=True)
add_para(
    'A: Pinia is Vue\'s official state management library. We use it to store the authentication '
    'state (who is logged in) so that it can be accessed from any component in the app. '
    'Without Pinia, we would need to pass the user data manually between components, which '
    'is complex and error-prone. Pinia provides a central store that any component can read from.'
)

# Q10
add_para('Q: What is the Vite proxy and why do you need it?', bold=True)
add_para(
    'A: During development, our Vue frontend runs on port 5173 and our PHP backend on port 8000. '
    'The Vite proxy forwards any request starting with /api to http://localhost:8000. This means '
    'the frontend can call /api/jobs instead of http://localhost:8000/api/jobs. This avoids '
    'CORS issues during development and makes the code cleaner. In production, both would be '
    'served from the same domain.'
)

# Q11
add_para('Q: How does the cascade delete work in the forum?', bold=True)
add_para(
    'A: In our database schema, the comments and post_likes tables have foreign keys to the '
    'posts table with "ON DELETE CASCADE". This means when a post is deleted, MySQL automatically '
    'deletes all related comments and likes. We do not need to write separate DELETE queries '
    'for comments and likes - the database handles it automatically.'
)

# Q12
add_para('Q: What HTTP status codes do you use and why?', bold=True)
add_para(
    'A: We use 200 for successful GET/PUT/DELETE, 201 for successful POST (creation), '
    '400 for bad input (validation errors), 401 for unauthorized (not logged in or bad token), '
    '403 for forbidden (wrong role), 404 for not found, 409 for conflicts (duplicates), '
    'and 500 for server errors. These are standard HTTP status codes that tell the frontend '
    'exactly what happened.'
)

# Q13
add_para('Q: How do you handle errors on the frontend?', bold=True)
add_para(
    'A: We use try/catch blocks in every Vue view that makes API calls. On error, we show '
    'a toast notification (red for errors) with the error message from the backend. '
    'The Axios response interceptor automatically handles 401 errors by clearing the token '
    'and redirecting to login. For form validation, we show inline error messages under each field.'
)

# Q14
add_para('Q: What is the upsert pattern in the profile controller?', bold=True)
add_para(
    'A: Upsert means "update or insert." We use MySQL\'s "INSERT ... ON DUPLICATE KEY UPDATE" '
    'syntax. The student_profiles table has a unique constraint on user_id. When a student '
    'saves their profile, if a profile already exists for their user_id, it updates the '
    'existing record. If not, it creates a new one. This means the same API endpoint '
    '(PUT /api/profile) handles both creating and updating.'
)

# Q15
add_para('Q: Why do you validate on both frontend and backend?', bold=True)
add_para(
    'A: Frontend validation provides instant feedback to users (no network round-trip). '
    'However, it can be bypassed - someone could send requests directly to the API using '
    'tools like Postman or curl. Backend validation is the real security barrier that '
    'cannot be bypassed. We do both for good user experience AND good security.'
)

# Q16
add_para('Q: What is the JSON response envelope and why do you use it?', bold=True)
add_para(
    'A: Every API response follows the same format: { success: boolean, message: string, '
    'data: object, errors: object }. This is called a response envelope. It makes it easy '
    'for the frontend to check if a request succeeded (check success field), show messages '
    'to the user (message field), access the data (data field), or show field-level errors '
    '(errors field). Consistency makes the frontend code simpler.'
)

# Q17
add_para('Q: How does the optimistic update work in ManageApplications?', bold=True)
add_para(
    'A: When an admin changes an application status, the frontend immediately updates the UI '
    'to show the new status (optimistic - assuming it will succeed). Then it sends the API '
    'request. If the request fails, the frontend rolls back the UI to the previous status '
    'and shows an error toast. This makes the app feel faster because the user sees the '
    'change immediately without waiting for the server.'
)

# Q18
add_para('Q: What is the debounce in the search box?', bold=True)
add_para(
    'A: In BrowseJobs.vue, the search input has a 400ms debounce. This means the search '
    'filter is not applied until the user stops typing for 400ms. Without debounce, every '
    'keystroke would trigger a filter, which is wasteful and can make the UI laggy. '
    'Debounce makes the search feel smooth and reduces unnecessary computation.'
)

add_page_break()

# ============================================================================
# 17. RUBRIC MAPPING
# ============================================================================

add_heading('17. Rubric Mapping (How We Meet Each Criterion)', level=1)

add_para('Here is how our project meets each criterion from the Demo Rubric (Section 8 of the project brief):')

add_table(
    ['Rubric Criterion', 'Weight', 'How We Meet It'],
    [
        ['SPA frontend quality', '15%',
         'Vue.js 3 SPA with Vue Router. Smooth page transitions with fade animation. '
         'Responsive design with Tailwind CSS (mobile hamburger menu, grid layouts). '
         'Role-based navigation. Loading spinners, toast notifications, confirm dialogs.'],
        ['CRUD completeness', '20%',
         'Full CRUD on jobs (GET, POST, PUT, DELETE). Full CRUD on companies. '
         'Applications: GET, POST, PUT (status), DELETE. Forum: full post CRUD + comments. '
         'Interviews: slot CRUD + booking create/update + evaluation.'],
        ['REST API implementation', '15%',
         'Slim 4 framework with proper routing. JSON responses with consistent envelope. '
         'Correct HTTP status codes (200, 201, 400, 401, 403, 404, 409, 500). '
         'Route parameters, query parameters, JSON body parsing.'],
        ['Database design and use', '10%',
         '10 related tables with foreign keys. ENUM types for roles, statuses, job types. '
         'Unique constraints (email, matric_no, job+user for applications, post+user for likes). '
         'Cascade delete on forum posts. Realistic seed data with Malaysian context.'],
        ['JWT and protected access', '15%',
         'JWT issued on login/register with HS256, 24h expiry. JwtMiddleware on every protected route. '
         'Role-based access control (student/admin/superadmin). Token stored in localStorage. '
         'Auto-logout on 401. Route guards on frontend.'],
        ['Validation and security practice', '10%',
         'Frontend validation on all forms. Backend validation on all endpoints. '
         'PDO prepared statements (SQL injection prevention). Bcrypt password hashing. '
         'CORS restrictions. JWT secret in .env. Ownership checks. Soft delete for jobs.'],
        ['Presentation and technical explanation', '15%',
         'This document. Clear architecture (frontend -> Vite proxy -> PHP backend -> MySQL). '
         'Module ownership clearly defined. Each team member can explain their modules. '
         'Code is well-commented and organized.'],
    ]
)

add_para('')
add_para('Required deliverables checklist:', bold=True)

add_table(
    ['Deliverable', 'Status', 'Location'],
    [
        ['Source code repository link', 'Done', 'github.com/saminsarwat007/webtechnologyproject'],
        ['Visible group contribution history', 'Done', 'Git commits + git-schedule/ folder with per-member plans'],
        ['Database schema/seed file', 'Done', 'database/schema.sql + database/seed.sql'],
        ['README with setup instructions', 'Done', 'README.md in project root'],
        ['SPA with Vue.js', 'Done', 'frontend/ folder with Vue 3 + Vite + Pinia + Vue Router'],
        ['REST API with GET/POST/PUT/DELETE', 'Done', 'backend/ folder with Slim 4, all 4 HTTP methods'],
        ['3+ related database tables', 'Done', '10 tables with foreign key relationships'],
        ['JWT authentication flow', 'Done', 'AuthController + JwtMiddleware + auth.js store'],
        ['Input validation (frontend + backend)', 'Done', 'Vue form validation + PHP controller validation'],
        ['Error handling with HTTP status codes', 'Done', 'Global error handler + proper status codes'],
    ]
)

add_page_break()

# ============================================================================
# APPENDIX: Quick Reference
# ============================================================================

add_heading('Appendix: Quick Reference Card', level=1)

add_para('Tech Stack:', bold=True)
add_bullet('Frontend: Vue 3 + Vite + Pinia + Vue Router + Tailwind CSS + Axios + Chart.js')
add_bullet('Backend: PHP + Slim 4 + firebase/php-jwt + vlucas/phpdotenv')
add_bullet('Database: MySQL with PDO')

add_para('')
add_para('Key Files:', bold=True)
add_bullet('Backend entry: backend/public/index.php (all routes)')
add_bullet('Database config: backend/src/config/database.php')
add_bullet('JWT middleware: backend/src/Middleware/JwtMiddleware.php')
add_bullet('Frontend entry: frontend/src/main.js')
add_bullet('Router: frontend/src/router/index.js')
add_bullet('Auth store: frontend/src/stores/auth.js')
add_bullet('API service: frontend/src/services/api.js')
add_bullet('Database schema: database/schema.sql')
add_bullet('Database seed: database/seed.sql')

add_para('')
add_para('Login Credentials (password: Password123!):', bold=True)
add_bullet('Super Admin: superadmin@careerbridge.my')
add_bullet('Admin: farizal@careerbridge.my')
add_bullet('Student: samin@student.utm.my')
add_bullet('Student: areeb@student.utm.my')

add_para('')
add_para('Server URLs (when running locally):', bold=True)
add_bullet('Frontend: http://localhost:5173')
add_bullet('Backend API: http://localhost:8000')
add_bullet('API health check: http://localhost:8000/')

add_para('')
add_para('Module Ownership:', bold=True)
add_bullet('M1 Auth + M6 Admin: Areeb (A22EC4041)')
add_bullet('M2 Jobs + M3 Applications: Samin (A22EC4040)')
add_bullet('M4 Companies + M5 Analytics: Mariam (A22EC4034)')
add_bullet('M7 Forum + M8 Interviews: Monika (A22EC4045)')

# ---- Save ----

output_path = os.path.join(os.path.dirname(__file__), 'CareerBridge_Codebase_Explanation.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Total paragraphs: {len(doc.paragraphs)}')
