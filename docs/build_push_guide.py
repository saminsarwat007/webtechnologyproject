"""
Build  docs/Push_Guide.docx  — a printable, week-by-week guide for every team
member describing exactly how to push their commits to GitHub.

It reads the source-of-truth markdown files under  git-schedule/  so the .docx
stays consistent with the per-member schedules.

Run:   docs/.venv/bin/python docs/build_push_guide.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# --- Paths ------------------------------------------------------------------

ROOT = Path(__file__).resolve().parent.parent
SCHEDULE_DIR = ROOT / "git-schedule"
OUTPUT = ROOT / "docs" / "Push_Guide.docx"

MEMBERS = [
    ("Mohammad Areeb",  "areeb",  "M1 Auth & Access Control · M6 System Admin"),
    ("Samin Sarwat",    "samin",  "M2 Job & Internship · M3 Application Tracking"),
    ("Mariam Hanif",    "mariam", "M4 Companies · M5 Reporting & Analytics"),
    ("Monika Zelenkov", "monika", "M7 Forum & Discussion · M8 Label Management"),
]

WEEKS = ["week6", "week7", "week8", "week9", "week10", "week11", "week12-14"]

REPO_URL = "https://github.com/azizah-utm/group-project-champion"

# --- Colours / sizing -------------------------------------------------------

NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT = RGBColor(0xB8, 0x6B, 0x00)   # warm gold for member names
GREY   = RGBColor(0x6C, 0x75, 0x7D)
BLACK  = RGBColor(0x1B, 0x1B, 0x1B)
CODE_BG = "F2F4F7"                    # light grey hex for code-block shading
NOTE_BG = "FFF7E6"                    # warm cream for callouts

# --- Low-level helpers ------------------------------------------------------

def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def heading(doc: Document, text: str, level: int = 1,
            color: RGBColor = NAVY, size: int | None = None,
            align: int = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.size = Pt(size if size else {1: 22, 2: 16, 3: 13}.get(level, 11))
    run.font.color.rgb = color
    if level in (1, 2, 3):
        p.style = doc.styles[f"Heading {level}"]


def para(doc: Document, text: str, size: int = 11,
         bold: bool = False, italic: bool = False,
         color: RGBColor = BLACK) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


# Inline-formatted paragraph: supports **bold**, *italic* and `code` segments.
INLINE_RX = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")

def rich_para(doc: Document, text: str, size: int = 11,
              left_indent: float | None = None) -> None:
    p = doc.add_paragraph()
    if left_indent is not None:
        p.paragraph_format.left_indent = Cm(left_indent)
    parts = INLINE_RX.split(text)
    for part in parts:
        if not part:
            continue
        run = p.add_run()
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.color.rgb = BLACK
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.font.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(size - 1)
            run.font.color.rgb = NAVY
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.font.italic = True
        else:
            run.text = part


def bullet(doc: Document, text: str, size: int = 11, indent: float = 0.5) -> None:
    # Render a bullet by manually prepending "•"; this avoids brittle docx list styles.
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    bullet_run = p.add_run("•  ")
    bullet_run.font.name = "Calibri"
    bullet_run.font.size = Pt(size)
    bullet_run.font.bold = True
    bullet_run.font.color.rgb = NAVY
    # rest of the line via inline parser
    parts = INLINE_RX.split(text)
    for part in parts:
        if not part:
            continue
        run = p.add_run()
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.color.rgb = BLACK
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]; run.font.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]; run.font.name = "Consolas"
            run.font.size = Pt(size - 1); run.font.color.rgb = NAVY
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]; run.font.italic = True
        else:
            run.text = part


def code_block(doc: Document, lines: list[str]) -> None:
    """Render a fenced code block as a single shaded cell with a monospaced font."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(16)
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(16)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_bg(cell, CODE_BG)
    # Clear default paragraph
    cell.paragraphs[0].text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = BLACK
    # tiny spacer after the block
    doc.add_paragraph()


def callout(doc: Document, text: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(16)
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(16)
    set_cell_bg(cell, NOTE_BG)
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    parts = INLINE_RX.split(text)
    for part in parts:
        if not part:
            continue
        run = p.add_run()
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]; run.font.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]; run.font.name = "Consolas"
            run.font.size = Pt(9.5); run.font.color.rgb = NAVY
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]; run.font.italic = True
        else:
            run.text = part
    doc.add_paragraph()


# --- Markdown parser (just enough for our schedule files) -------------------

def render_markdown(doc: Document, md_text: str, *, base_heading: int = 2) -> None:
    """Render a small subset of markdown into the docx document.

    Supports: # H1, ## H2, ### H3, **bold**, *italic*, `code`, fenced ```code```
    blocks, "- " bullets, and "> blockquote" callouts. Anything else is treated
    as a regular paragraph.

    `base_heading` shifts all H1 in the markdown to that docx heading level
    (so a per-week file's "# Samin — Week 6" can render at H3 inside the
    member chapter).
    """
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]
        # Fenced code
        if line.strip().startswith("```"):
            if in_code:
                code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Horizontal rule → spacer
        if line.strip() == "---":
            para(doc, "")
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            heading(doc, line[4:].strip(), level=min(base_heading + 2, 3))
        elif line.startswith("## "):
            heading(doc, line[3:].strip(), level=min(base_heading + 1, 3))
        elif line.startswith("# "):
            heading(doc, line[2:].strip(), level=base_heading)
        elif line.startswith("> "):
            callout(doc, line[2:].strip())
        elif line.startswith("- "):
            bullet(doc, line[2:].strip())
        elif line.strip() == "":
            # Soft spacer only when previous line was content
            pass
        else:
            rich_para(doc, line.rstrip())
        i += 1


# --- Document sections ------------------------------------------------------

def add_cover(doc: Document) -> None:
    # Vertical-centered title block
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CareerBridge")
    r.font.name = "Calibri"; r.font.size = Pt(44); r.font.bold = True
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Weekly Push Guide for the Team")
    r.font.name = "Calibri"; r.font.size = Pt(20); r.font.color.rgb = GREY

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Step-by-step commands every member runs each week")
    r.font.name = "Calibri"; r.font.size = Pt(13); r.font.italic = True
    r.font.color.rgb = GREY

    for _ in range(4):
        doc.add_paragraph()

    # Members box
    tbl = doc.add_table(rows=len(MEMBERS) + 1, cols=2)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(6)
    tbl.columns[1].width = Cm(10)

    hdr = tbl.rows[0].cells
    for c in hdr: set_cell_bg(c, "1E3A5F")
    for cell, text in zip(hdr, ("Member", "Modules Owned")):
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(text)
        run.font.name = "Calibri"; run.font.bold = True
        run.font.size = Pt(11); run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row, (name, _, modules) in zip(tbl.rows[1:], MEMBERS):
        c0, c1 = row.cells
        c0.width = Cm(6); c1.width = Cm(10)
        c0.paragraphs[0].text = ""
        r0 = c0.paragraphs[0].add_run(name)
        r0.font.name = "Calibri"; r0.font.size = Pt(11); r0.font.bold = True
        c1.paragraphs[0].text = ""
        r1 = c1.paragraphs[0].add_run(modules)
        r1.font.name = "Calibri"; r1.font.size = Pt(10.5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Repository: {REPO_URL}")
    r.font.name = "Consolas"; r.font.size = Pt(10.5); r.font.color.rgb = NAVY

    add_page_break(doc)


def add_intro(doc: Document) -> None:
    heading(doc, "How this guide works", level=1)
    rich_para(doc,
        "This document is split into four sections. Read sections 1–2 once, "
        "then jump to your own chapter and follow the commands week by week.")

    bullet(doc, "**Section 1 — One-time setup.** Run *once* on your laptop. "
                "After this you'll never run these commands again.")
    bullet(doc, "**Section 2 — Weekly rhythm.** The Day 1 / Day 2 / Day 3 pattern "
                "every member follows every week.")
    bullet(doc, "**Section 3 — Per-member chapters.** Four chapters, one per teammate, "
                "with 7 weekly sub-sections each. Pure copy-paste.")
    bullet(doc, "**Section 4 — Recovery & PR etiquette.** What to do when something "
                "goes wrong.")

    callout(doc,
        "Every command shown in a grey box is meant to be **copy-pasted into your "
        "terminal exactly as written**. Lines that begin with `#` are comments — "
        "you can paste them too; the shell will ignore them.")

    add_page_break(doc)


def add_one_time_setup(doc: Document) -> None:
    heading(doc, "Section 1 — One-time setup", level=1)
    rich_para(doc, "Do this *once* on the laptop you'll be committing from.")

    heading(doc, "1.1 Install Git and identify yourself", level=2)
    rich_para(doc, "Open the macOS / Linux terminal (or Git Bash on Windows) and run:")
    code_block(doc, [
        "git --version              # should print something like git version 2.x",
        "",
        "# Replace these two lines with YOUR name and the email on YOUR GitHub account",
        'git config --global user.name  "Mohammad Areeb"',
        'git config --global user.email "areeb@example.com"',
        "",
        "git config --global init.defaultBranch main",
        "git config --global pull.rebase    false",
    ])

    heading(doc, "1.2 Authenticate with GitHub", level=2)
    rich_para(doc,
        "Pick ONE of these. Option A is easier on the first try; Option B is "
        "better for long-term work.")

    heading(doc, "Option A — Personal Access Token (recommended for first push)", level=3)
    bullet(doc, "Go to GitHub → **Settings → Developer settings → Personal access tokens "
                "→ Tokens (classic) → Generate new token**.")
    bullet(doc, "Tick the `repo` scope. Set expiry to 90 days.")
    bullet(doc, "Copy the token *immediately*. GitHub will never show it again.")
    bullet(doc, "The first time you `git push`, paste the token when prompted for "
                "**Password**.")

    heading(doc, "Option B — SSH key", level=3)
    code_block(doc, [
        'ssh-keygen -t ed25519 -C "you@example.com"',
        "cat ~/.ssh/id_ed25519.pub",
    ])
    rich_para(doc,
        "Copy the printed key into GitHub → **Settings → SSH and GPG keys → "
        "New SSH key**.")

    heading(doc, "1.3 Clone the repository", level=2)
    code_block(doc, [
        "# Use the HTTPS URL if you picked Option A above:",
        f"git clone {REPO_URL}.git",
        "",
        "# …or the SSH URL if you picked Option B:",
        "git clone git@github.com:azizah-utm/group-project-champion.git",
        "",
        "cd group-project-champion",
    ])

    heading(doc, "1.4 Smoke-test that you can push", level=2)
    code_block(doc, [
        "git checkout -b sandbox/your-name-test",
        "echo test > _delete_me.txt",
        "git add _delete_me.txt",
        'git commit -m "test: verify I can push"',
        "git push -u origin sandbox/your-name-test",
    ])
    rich_para(doc, "If that succeeds, delete the sandbox branch locally and on GitHub:")
    code_block(doc, [
        "git checkout main",
        "git branch -D sandbox/your-name-test",
        "git push origin --delete sandbox/your-name-test",
        "rm _delete_me.txt",
    ])

    add_page_break(doc)


def add_weekly_rhythm(doc: Document) -> None:
    heading(doc, "Section 2 — The weekly rhythm", level=1)
    rich_para(doc,
        "Every week, every member does this same three-step dance. The exact "
        "branch name, files to touch, and commit messages come from your own "
        "weekly chapter further down in this document.")

    heading(doc, "Day 1 (Monday) — start the week's branch", level=2)
    code_block(doc, [
        "git checkout main",
        "git pull origin main",
        "git checkout -b feature/<branch-name-from-your-weekN>",
        "",
        "# …make Day 1 edits…",
        "git status",
        "git add <only the files listed for Day 1>",
        'git commit -m "<exact commit message from your weekN>"',
        "git push -u origin feature/<branch-name>",
    ])

    heading(doc, "Day 2 (Wednesday) — mid-week commit", level=2)
    code_block(doc, [
        "git pull origin feature/<branch-name>     # in case you pushed from elsewhere",
        "# …make Day 2 edits…",
        "git add <Day-2 files>",
        'git commit -m "<exact commit message from your weekN>"',
        "git push origin feature/<branch-name>",
    ])

    heading(doc, "Day 3 (Friday) — final commit + Pull Request", level=2)
    code_block(doc, [
        "# …make Day 3 edits…",
        "git add <Day-3 files>",
        'git commit -m "<exact commit message from your weekN>"',
        "git push origin feature/<branch-name>",
    ])
    rich_para(doc,
        "Then open GitHub in the browser, click **Compare & pull request**, "
        "set base = `main`, compare = your branch, add the other three members "
        "as reviewers, and create the PR.")

    heading(doc, "After the PR is merged — clean up", level=2)
    code_block(doc, [
        "git checkout main",
        "git pull origin main",
        "git branch  -d feature/<branch-name>           # delete local copy",
        "git push origin --delete feature/<branch-name> # delete remote copy",
    ])

    add_page_break(doc)


def add_member_chapter(doc: Document, name: str, slug: str, modules: str) -> None:
    # Chapter title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(name.upper())
    r.font.name = "Calibri"; r.font.size = Pt(28); r.font.bold = True
    r.font.color.rgb = ACCENT
    p.style = doc.styles["Heading 1"]

    rich_para(doc, f"**Modules owned:** {modules}")
    para(doc, "")

    # Render every weekly markdown file for this member
    for w in WEEKS:
        path = SCHEDULE_DIR / slug / f"{w}.md"
        if not path.exists():
            rich_para(doc, f"*Missing file: {path.relative_to(ROOT)}*")
            continue
        md = path.read_text(encoding="utf-8")
        # base_heading=2 means "# Title" in the file becomes H2 in the docx
        render_markdown(doc, md, base_heading=2)
        # small spacer between weeks
        para(doc, "")

    add_page_break(doc)


def add_recovery(doc: Document) -> None:
    heading(doc, "Section 4 — Recovery & PR etiquette", level=1)
    rich_para(doc,
        "Almost every git emergency is fixable in one command — but the wrong "
        "command can lose a day's work. Before typing anything destructive, "
        "screenshot the error and ask the group chat.")

    heading(doc, "4.1 I committed to `main` by accident", level=2)
    code_block(doc, [
        "git branch feature/<correct-branch-name>",
        "git reset --hard origin/main             # ⚠ destroys local main changes",
        "git checkout feature/<correct-branch-name>",
        "git push -u origin feature/<correct-branch-name>",
    ])

    heading(doc, "4.2 I committed the wrong files", level=2)
    code_block(doc, [
        "git reset --soft HEAD~1                  # undo commit, keep changes staged",
        "git restore --staged <wrong-file>",
        "git add <correct-files>",
        'git commit -m "<correct message>"',
        "git push --force-with-lease origin feature/<branch>",
    ])
    callout(doc,
        "**Never `--force` on `main`.** `--force-with-lease` is safer because "
        "it refuses to overwrite if a teammate has pushed in the meantime.")

    heading(doc, "4.3 `git push` rejected because remote is ahead", level=2)
    code_block(doc, [
        "git pull origin feature/<branch>",
        "# resolve any merge conflicts (see 4.4)",
        "git push origin feature/<branch>",
    ])

    heading(doc, "4.4 Merge conflict during pull", level=2)
    rich_para(doc, "Git marks each conflict like this inside the file:")
    code_block(doc, [
        "<<<<<<< HEAD",
        "your version",
        "=======",
        "their version",
        ">>>>>>> origin/feature/<branch>",
    ])
    bullet(doc, "Decide which side wins (or merge by hand).")
    bullet(doc, "**Delete the `<<<<<<<`, `=======`, and `>>>>>>>` marker lines.**")
    bullet(doc, "Save the file.")
    code_block(doc, [
        "git add <resolved files>",
        "git commit                # git fills in a merge message",
        "git push origin feature/<branch>",
    ])
    rich_para(doc, "If you want to abandon the merge entirely:")
    code_block(doc, ["git merge --abort"])

    heading(doc, "4.5 I accidentally pushed something secret (.env, password)", level=2)
    rich_para(doc, "Tell the group chat **immediately**, then:")
    code_block(doc, [
        "git rm --cached backend/.env",
        'echo "backend/.env" >> .gitignore',
        "git add .gitignore",
        'git commit -m "chore: remove leaked .env from history"',
        "git push --force-with-lease origin feature/<branch>",
    ])
    rich_para(doc,
        "Then **rotate the secret** — change the JWT secret and the DB password "
        "even if the leak only lived on a feature branch.")

    heading(doc, "4.6 PR etiquette", level=2)
    bullet(doc, "Title and body must reference your module "
                "(e.g. *Module 7: forum post detail page*).")
    bullet(doc, "Add the other three members as reviewers.")
    bullet(doc, "Never click **Merge** on your own PR until at least one reviewer approves.")
    bullet(doc, "When you review someone else's PR: pull their branch, run "
                "`composer install` / `npm install`, smoke-test it, then click **Approve**.")

    heading(doc, "4.7 The Forbidden List", level=2)
    bullet(doc, "❌ `git push --force` on `main`.")
    bullet(doc, "❌ Deleting the `main` branch.")
    bullet(doc, "❌ Committing real `backend/.env` (the example one is fine).")
    bullet(doc, "❌ Committing `node_modules/` or `vendor/`.")
    bullet(doc, "❌ Editing another member's controller without telling them.")
    bullet(doc, "❌ Merging your own PR with zero reviews.")


# --- Main -------------------------------------------------------------------

def main() -> None:
    doc = Document()
    # Page setup: A4 with 2 cm margins
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = section.right_margin = Cm(2.0)
        section.top_margin = section.bottom_margin = Cm(2.0)

    add_cover(doc)
    add_intro(doc)
    add_one_time_setup(doc)
    add_weekly_rhythm(doc)

    heading(doc, "Section 3 — Per-member chapters", level=1)
    rich_para(doc,
        "Each member has their own chapter below containing all seven weeks "
        "(Week 6 through Week 12-14). Skip straight to your own name.")
    add_page_break(doc)

    for name, slug, modules in MEMBERS:
        add_member_chapter(doc, name, slug, modules)

    add_recovery(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT.relative_to(ROOT)}  ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
