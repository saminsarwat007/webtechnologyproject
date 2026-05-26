"""
Build the CareerBridge Phase 1 Project Proposal as docs/Proposal.docx.

Run:  docs/.venv/bin/python docs/build_proposal.py

Requires  docs/diagrams/*.png  and  docs/mockups/*.png  to already exist
(run docs/build_diagrams.py first).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
DIAGRAMS = ROOT / "diagrams"
MOCKUPS = ROOT / "mockups"
OUTPUT = ROOT / "Proposal.docx"

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
GREY = RGBColor(0x6C, 0x75, 0x7D)
BLACK = RGBColor(0x1B, 0x1B, 0x1B)


# --- Helpers ----------------------------------------------------------------

def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def heading(doc: Document, text: str, level: int = 1,
            color: RGBColor = NAVY, size: int = 16,
            align: int = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    # Tag as heading for Word's navigation pane
    p.style = doc.styles[f"Heading {level}"] if level in (1, 2, 3) else p.style


def para(doc: Document, text: str, size: int = 11,
         bold: bool = False, italic: bool = False,
         color: RGBColor = BLACK, align: int = WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def bullet(doc: Document, text: str, indent: float = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_image(doc: Document, path: Path, width_cm: float = 16.0,
              caption: str = "") -> None:
    doc.add_picture(str(path), width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = GREY


def make_table(doc: Document, headers: list[str], rows: list[list[str]],
               widths_cm: list[float] | None = None,
               header_bg: str = "1e3a5f") -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(hdr_cells[i], header_bg)
    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, value in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(10)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for row in t.rows:
                row.cells[i].width = Cm(w)


# --- Build the document -----------------------------------------------------

def build() -> None:
    doc = Document()

    # ---- Default style tweak
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # =========================================================
    # COVER PAGE
    # =========================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n")
    # School
    heading(doc, "FACULTY OF COMPUTING", level=1, size=20,
            align=WD_ALIGN_PARAGRAPH.CENTER)
    heading(doc, "UNIVERSITI TEKNOLOGI MALAYSIA",
            level=2, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
            color=GREY)
    para(doc, " ")
    heading(doc, "SECJ3483 — WEB TECHNOLOGY", level=2, size=16,
            align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "GROUP PROJECT — SECTION XX",
         bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
    para(doc, " ")
    para(doc, " ")
    heading(doc, "Phase 1: Project Proposal", level=1, size=22,
            align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, " ")

    para(doc, "Project Title", bold=True, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
    heading(doc, "CAREERBRIDGE", level=2, size=22,
            align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Student Internship & Job Application Tracker",
         italic=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
    para(doc, " ")

    para(doc, "Team 1 — CHAMPION", bold=True, size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Team Motto: \u201cBEING CHAMPION\u201d",
         italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
    para(doc, " ")

    para(doc, "Lecturer:", bold=True, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "[ To be filled in ]",
         size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
    para(doc, " ")
    para(doc, "Prepared by:", bold=True, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, " ")

    make_table(doc,
               headers=["NO.", "NAME", "MATRICS NUMBER"],
               rows=[
                   ["1.", "Mohammad Areeb", "A22EC4035"],
                   ["2.", "Samin Sarwat",   "A22EC4040"],
                   ["3.", "Mariam Hanif",   "A22EC4034"],
               ],
               widths_cm=[2.0, 8.0, 5.0])

    para(doc, " ")
    para(doc, "Academic Session: 2025 / 2026",
         size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
    para(doc, f"Submission Date: 2 May 2026",
         size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)

    add_page_break(doc)

    # =========================================================
    # TABLE OF CONTENTS (manual)
    # =========================================================
    heading(doc, "Table of Contents", level=1, size=18)
    toc = [
        ("1. Problem Statement and User Focus", "3"),
        ("    1.1 Problem Statement", "3"),
        ("    1.2 Target Users", "3"),
        ("    1.3 Proposed Solution", "4"),
        ("2. Functional Scope and Modules", "5"),
        ("    2.1 System Modules", "5"),
        ("    2.2 Use Case Diagram", "6"),
        ("    2.3 Functional Requirements", "7"),
        ("3. Technical Design", "8"),
        ("    3.1 System Architecture", "8"),
        ("    3.2 Technology Stack", "9"),
        ("    3.3 Database Design (ER Diagram)", "9"),
        ("    3.4 API Design", "10"),
        ("    3.5 Security", "11"),
        ("4. Interface Planning (Mockups)", "12"),
        ("    4.1 Login Screen", "12"),
        ("    4.2 Browse Jobs", "12"),
        ("    4.3 Student Dashboard", "13"),
        ("    4.4 Admin Analytics Dashboard", "13"),
        ("5. Team Planning", "14"),
        ("    5.1 Module Ownership", "14"),
        ("    5.2 9-Week Schedule", "14"),
        ("    5.3 Collaboration Workflow", "15"),
    ]
    for title, page in toc:
        p = doc.add_paragraph()
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(16.0), WD_ALIGN_PARAGRAPH.RIGHT,
                               leader=4)  # 4 == dotted leader
        run = p.add_run(title)
        run.font.size = Pt(11)
        run2 = p.add_run("\t" + page)
        run2.font.size = Pt(11)

    add_page_break(doc)

    # =========================================================
    # SECTION 1 — Problem Statement and User Focus
    # =========================================================
    heading(doc, "1. Problem Statement and User Focus", level=1, size=18)

    heading(doc, "1.1 Problem Statement", level=2, size=13)
    para(doc,
         "University students in Malaysia routinely struggle to find relevant "
         "internship and graduate employment opportunities. Listings are "
         "scattered across LinkedIn, company career pages, lecturer emails, "
         "Facebook groups and university notice boards. Students frequently "
         "miss deadlines, apply to roles they do not qualify for, and lose "
         "track of which companies they have already contacted. At the same "
         "time, university career-services staff lack a consolidated system "
         "for publishing approved opportunities, reviewing student "
         "applications, and reporting placement outcomes to the faculty.")
    para(doc,
         "The core problems this project addresses are:")
    bullet(doc, "Fragmented job sources — students waste time hunting across "
                "multiple platforms.")
    bullet(doc, "No central application tracker — students forget which jobs "
                "they applied for and their current status.")
    bullet(doc, "No governance on postings — fraudulent or irrelevant listings "
                "can reach students through informal channels.")
    bullet(doc, "No analytics — career-services admins cannot answer basic "
                "questions like \u201chow many students applied this week?\u201d")

    heading(doc, "1.2 Target Users", level=2, size=13)
    make_table(doc,
        headers=["User Role", "Primary Goals", "Key Actions"],
        rows=[
            ["Student",
             "Find and secure internships / jobs that match their skills.",
             "Register, build profile, browse and filter listings, apply with a "
             "cover letter, track application status, withdraw pending applications."],
            ["Admin (Career Services)",
             "Publish vetted opportunities and moderate applications.",
             "Create / edit / remove job postings, manage partner companies, "
             "update application statuses, view analytics."],
            ["Super Admin",
             "Oversee the platform and manage user accounts.",
             "Manage users, promote admins, access all analytics, moderate "
             "across the entire system."],
        ],
        widths_cm=[3.5, 5.0, 8.0])
    para(doc, " ")

    heading(doc, "1.3 Proposed Solution", level=2, size=13)
    para(doc,
         "CareerBridge is a full-stack web application that consolidates the "
         "internship and job search into a single platform tailored to "
         "university career-services workflows. Students receive a clean, "
         "filterable feed of active opportunities, a one-click apply flow "
         "with cover-letter support, and a persistent dashboard that shows "
         "the status of every application they have ever submitted. Admins "
         "receive CRUD tools for jobs, applications and companies plus a "
         "dashboard with live platform metrics. Role-based access control "
         "enforced by JWT ensures that each user sees only what is "
         "appropriate to their role.")

    add_page_break(doc)

    # =========================================================
    # SECTION 2 — Functional Scope and Modules
    # =========================================================
    heading(doc, "2. Functional Scope and Modules", level=1, size=18)

    heading(doc, "2.1 System Modules", level=2, size=13)
    para(doc,
         "The system is decomposed into six functional modules. Each module "
         "is owned end-to-end by one team member who implements both the "
         "backend CRUD endpoints and the matching Vue views.")
    make_table(doc,
        headers=["#", "Module", "Owner", "Scope Summary"],
        rows=[
            ["M1", "User Authentication & Access Control", "Areeb",
             "Registration, login, JWT issuance, role-based route guards, "
             "password hashing with bcrypt."],
            ["M2", "Job & Internship Management", "Samin",
             "Full CRUD on job listings with search / filter by type and "
             "company; soft-delete via is_active flag."],
            ["M3", "Application Tracking System", "Samin",
             "Full CRUD on applications: apply, track status, admin moderation, "
             "student withdrawal (own + pending only)."],
            ["M4", "Company & External Source Management", "Mariam",
             "Full CRUD on partner companies with FK guard preventing deletion "
             "while jobs are still linked."],
            ["M5", "Reporting & Analytics", "Mariam",
             "Aggregate dashboards for admins (platform-wide) and students "
             "(personal); student profile upsert for dashboard data."],
            ["M6", "System Administration & Notifications", "Areeb",
             "Super-admin user directory, global error handler, toast "
             "notifications, confirm dialogs."],
        ],
        widths_cm=[1.0, 4.5, 2.5, 8.5])
    para(doc, " ")

    heading(doc, "2.2 Use Case Diagram", level=2, size=13)
    para(doc,
         "The diagram below summarises the primary interactions between the "
         "three actors (Student, Admin, Super Admin) and the CareerBridge "
         "system. \u00abinclude\u00bb relationships indicate that one use "
         "case always triggers another (e.g. Apply requires prior Login). "
         "\u00abextend\u00bb relationships indicate an optional extension "
         "available only to a subset of actors (Super Admin can Manage Users).")
    add_image(doc, DIAGRAMS / "01_use_case_diagram.png", width_cm=16,
              caption="Figure 2.1 — CareerBridge Use Case Diagram")

    add_page_break(doc)

    heading(doc, "2.3 Functional Requirements", level=2, size=13)
    para(doc,
         "The following table lists the functional requirements by module. "
         "Each requirement maps directly to at least one CRUD endpoint in "
         "Section 3.4.")
    make_table(doc,
        headers=["Req. ID", "Module", "Requirement"],
        rows=[
            ["FR-1.1", "M1", "A visitor shall be able to register as a student or admin with full name, email, and password (min 8 chars)."],
            ["FR-1.2", "M1", "A registered user shall be able to log in and receive a JWT valid for 24 hours."],
            ["FR-1.3", "M1", "The system shall restrict protected endpoints by role (student / admin / superadmin)."],
            ["FR-2.1", "M2", "An admin shall be able to create a job with title, type, description, requirements and deadline."],
            ["FR-2.2", "M2", "A student shall be able to browse active jobs, filtered by keyword, type, and company."],
            ["FR-2.3", "M2", "An admin shall be able to edit or soft-delete any job."],
            ["FR-3.1", "M3", "A student shall be able to apply to a job with an optional cover letter (max 5000 chars)."],
            ["FR-3.2", "M3", "The system shall prevent a student from applying to the same job twice."],
            ["FR-3.3", "M3", "A student shall be able to withdraw a pending application."],
            ["FR-3.4", "M3", "An admin shall be able to move an application through states pending → reviewed → accepted / rejected."],
            ["FR-4.1", "M4", "An admin shall be able to create, edit and delete company records."],
            ["FR-4.2", "M4", "The system shall reject deletion of a company while jobs reference it."],
            ["FR-5.1", "M5", "An admin shall see aggregate counts for jobs, applications, and applications-this-week."],
            ["FR-5.2", "M5", "A student shall see personal statistics and a list of their recent applications."],
            ["FR-5.3", "M5", "A student shall be able to upsert their profile (matric, programme, CGPA, skills, resume text)."],
            ["FR-6.1", "M6", "A super-admin shall be able to list all users with their roles."],
            ["FR-6.2", "M6", "The system shall surface errors and success events via non-blocking toast notifications."],
            ["FR-6.3", "M6", "Destructive actions (delete job / company / application) shall require an explicit confirmation dialog."],
        ],
        widths_cm=[1.8, 1.4, 12.8])

    add_page_break(doc)

    # =========================================================
    # SECTION 3 — Technical Design
    # =========================================================
    heading(doc, "3. Technical Design", level=1, size=18)

    heading(doc, "3.1 System Architecture", level=2, size=13)
    para(doc,
         "CareerBridge follows a classic three-tier architecture. The "
         "Presentation tier is a single-page Vue 3 application served by "
         "Vite during development and by a static host in production. It "
         "communicates exclusively over HTTPS JSON with the Application "
         "tier, a stateless Slim 4 REST API written in PHP. The Data tier "
         "is a MySQL (or MariaDB) server accessed only through PDO prepared "
         "statements. This separation means each tier can scale, be "
         "replaced, or be re-implemented independently.")
    add_image(doc, DIAGRAMS / "02_architecture_diagram.png", width_cm=16,
              caption="Figure 3.1 — Three-tier system architecture")

    heading(doc, "3.2 Technology Stack", level=2, size=13)
    make_table(doc,
        headers=["Layer", "Technology", "Justification"],
        rows=[
            ["Frontend framework", "Vue 3 + Vite",
             "Composition API is beginner-friendly and Vite offers instant HMR."],
            ["State management", "Pinia",
             "Official Vue 3 store with TypeScript support and devtools."],
            ["Routing", "Vue Router 4",
             "Native navigation guards for role-based access control."],
            ["Styling", "Tailwind CSS",
             "Utility-first, consistent spacing, small production bundle."],
            ["Charts", "Chart.js",
             "Mature, accessible library for dashboards."],
            ["HTTP", "Axios",
             "Request / response interceptors simplify JWT attachment and "
             "401 handling."],
            ["Backend framework", "Slim 4 (PHP)",
             "Tiny footprint, explicit routing, middleware pipeline."],
            ["Auth", "firebase/php-jwt",
             "Battle-tested HS256 JWT signing."],
            ["Database", "MySQL 5.7+ / MariaDB 10.3+",
             "Widely available on XAMPP / MAMP / Laragon for assessment demos."],
            ["DB access", "PDO (prepared statements)",
             "Eliminates SQL-injection surface; no ORM overhead."],
        ],
        widths_cm=[3.2, 4.0, 9.0])
    para(doc, " ")

    heading(doc, "3.3 Database Design (ER Diagram)", level=2, size=13)
    para(doc,
         "The schema contains five tables with referential integrity "
         "enforced by InnoDB foreign keys. A unique index on "
         "(job_id, user_id) in the applications table prevents duplicate "
         "applications at the database level, not just at the application "
         "level. ENUM types are used for finite domains (role, job type, "
         "application status) so invalid values cannot be inserted.")
    add_image(doc, DIAGRAMS / "03_er_diagram.png", width_cm=17,
              caption="Figure 3.2 — Entity-Relationship diagram")

    heading(doc, "3.4 API Design", level=2, size=13)
    para(doc,
         "All endpoints are JSON and follow a uniform envelope "
         "{ success, message, data?, errors? }. Protected routes require a "
         "Bearer JWT in the Authorization header. The table below lists "
         "every endpoint grouped by module.")

    make_table(doc,
        headers=["Method", "Path", "Auth", "Module"],
        rows=[
            ["POST", "/api/auth/register", "Public", "M1"],
            ["POST", "/api/auth/login", "Public", "M1"],
            ["GET",  "/api/jobs", "Public", "M2"],
            ["GET",  "/api/jobs/{id}", "Public", "M2"],
            ["POST", "/api/jobs", "admin, superadmin", "M2"],
            ["PUT",  "/api/jobs/{id}", "admin, superadmin", "M2"],
            ["DELETE", "/api/jobs/{id}", "admin, superadmin", "M2"],
            ["GET",  "/api/applications", "Any logged-in", "M3"],
            ["POST", "/api/applications", "student", "M3"],
            ["PUT",  "/api/applications/{id}/status", "admin, superadmin", "M3"],
            ["DELETE", "/api/applications/{id}", "student (own + pending)", "M3"],
            ["GET",  "/api/companies", "Any logged-in", "M4"],
            ["POST", "/api/companies", "admin, superadmin", "M4"],
            ["PUT",  "/api/companies/{id}", "admin, superadmin", "M4"],
            ["DELETE", "/api/companies/{id}", "admin, superadmin", "M4"],
            ["GET",  "/api/profile", "student", "M5"],
            ["PUT",  "/api/profile", "student", "M5"],
            ["GET",  "/api/admin/analytics", "admin, superadmin", "M5"],
            ["GET",  "/api/admin/users", "superadmin", "M6"],
        ],
        widths_cm=[1.8, 6.8, 5.0, 1.4])
    para(doc, " ")

    heading(doc, "3.5 Security", level=2, size=13)
    bullet(doc, "Passwords are stored using PASSWORD_BCRYPT (cost 12); "
                "plaintext is never persisted.")
    bullet(doc, "JWTs are signed with HS256 using a per-environment secret "
                "loaded from .env; tokens expire after 24 hours.")
    bullet(doc, "Every SQL query uses PDO prepared statements with named "
                "placeholders — no string concatenation with user input.")
    bullet(doc, "CORS is restricted to the configured frontend origin by "
                "default; wildcard mode is opt-in via CORS_ALLOW_ALL=1 for "
                "development only.")
    bullet(doc, "The Axios interceptor clears the token and redirects to "
                "/login on any 401 response.")
    bullet(doc, "Destructive frontend actions require an explicit confirm "
                "dialog before the DELETE request is sent.")

    add_page_break(doc)

    # =========================================================
    # SECTION 4 — Interface Planning (Mockups)
    # =========================================================
    heading(doc, "4. Interface Planning (Mockups)", level=1, size=18)
    para(doc,
         "The mockups below illustrate the four most-used screens in the "
         "application. Each mockup reflects the final UI direction: Tailwind "
         "utility classes for spacing, a navy-and-amber brand palette, and "
         "responsive grid layouts that collapse to a single column on mobile.")

    heading(doc, "4.1 Login Screen", level=2, size=13)
    para(doc,
         "Single centred card, minimal fields, success redirects to the "
         "role-appropriate dashboard. Failed attempts surface a toast "
         "rather than a full-page error.")
    add_image(doc, MOCKUPS / "01_login.png", width_cm=14,
              caption="Figure 4.1 — Login screen")

    heading(doc, "4.2 Browse Jobs (Student)", level=2, size=13)
    para(doc,
         "Filter row at top (debounced search, type dropdown, company "
         "dropdown) with a responsive card grid. Cards are keyboard-"
         "focusable and each exposes a single primary action.")
    add_image(doc, MOCKUPS / "02_browse_jobs.png", width_cm=16,
              caption="Figure 4.2 — Browse Jobs page")

    add_page_break(doc)

    heading(doc, "4.3 Student Dashboard", level=2, size=13)
    para(doc,
         "Four metric cards across the top summarise the student's current "
         "application pipeline. A recent-applications table gives quick "
         "context and colour-coded status badges for scannability.")
    add_image(doc, MOCKUPS / "03_student_dashboard.png", width_cm=16,
              caption="Figure 4.3 — Student Dashboard")

    heading(doc, "4.4 Admin Analytics Dashboard", level=2, size=13)
    para(doc,
         "Admins see platform-wide metrics and a Chart.js bar chart that "
         "breaks down applications by status. All numbers are fetched from "
         "a single /api/admin/analytics call.")
    add_image(doc, MOCKUPS / "04_admin_dashboard.png", width_cm=16,
              caption="Figure 4.4 — Admin Analytics Dashboard")

    add_page_break(doc)

    # =========================================================
    # SECTION 5 — Team Planning
    # =========================================================
    heading(doc, "5. Team Planning", level=1, size=18)

    heading(doc, "5.1 Module Ownership", level=2, size=13)
    para(doc,
         "Every member implements full CRUD for the modules they own, "
         "including both the backend endpoints and the matching Vue views. "
         "This vertical-slice split makes individual contribution auditable "
         "from the git log alone.")
    make_table(doc,
        headers=["Member", "Matric", "Modules", "CRUD Coverage"],
        rows=[
            ["Mohammad Areeb", "A22EC4035",
             "M1 Auth & Access Control · M6 System Admin & Notifications",
             "Create + Read on users (register / login), Read on users "
             "(admin directory), cross-cutting middleware."],
            ["Samin Sarwat", "A22EC4040",
             "M2 Job & Internship Mgmt · M3 Application Tracking",
             "Full CRUD on jobs AND full CRUD on applications — 8 endpoints."],
            ["Mariam Hanif", "A22EC4034",
             "M4 Company Mgmt · M5 Reporting & Analytics",
             "Full CRUD on companies, aggregate Read for analytics, "
             "Create + Update on student_profiles via INSERT … ON DUPLICATE "
             "KEY UPDATE."],
        ],
        widths_cm=[3.0, 2.2, 5.5, 5.5])
    para(doc, " ")

    heading(doc, "5.2 9-Week Schedule", level=2, size=13)
    para(doc,
         "Development runs across 9 teaching weeks (Wk6 → Wk14). Every "
         "member opens a feature branch on Day 1 of each week, commits "
         "three times (Day 1, Day 2, Day 3), and merges to main via Pull "
         "Request with a peer review at the end of the week.")
    add_image(doc, DIAGRAMS / "04_gantt_chart.png", width_cm=17,
              caption="Figure 5.1 — 9-week project Gantt chart")

    heading(doc, "5.3 Collaboration Workflow", level=2, size=13)
    numbered(doc, "Branch per week: feature/<module>-<phase> (e.g. "
                  "feature/jobs-api, feature/applications-frontend).")
    numbered(doc, "Conventional commit messages: feat / fix / chore / docs / "
                  "refactor with a module-scoped prefix.")
    numbered(doc, "Every Pull Request needs one peer approval before merging "
                  "to main.")
    numbered(doc, "Folder ownership is documented in "
                  "git-schedule/HOW_TO_PUSH.md so that members do not edit "
                  "each other's modules by accident.")
    numbered(doc, "Weekly status syncs every Friday: what shipped, what is "
                  "blocked, what is next.")

    heading(doc, "5.4 Deliverables", level=2, size=13)
    bullet(doc, "Phase 1 (Wk5): this proposal document.")
    bullet(doc, "Phase 2 (Wk9): mid-project demo with M1 + M2 + M3 "
                "working end-to-end.")
    bullet(doc, "Phase 3 (Wk14): final demo, tagged v1.0.0 release, and "
                "project report.")

    # =========================================================
    # Save
    # =========================================================
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT.relative_to(ROOT.parent)}")


if __name__ == "__main__":
    build()
